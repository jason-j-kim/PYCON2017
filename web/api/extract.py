# -*- coding: utf-8 -*-
"""Vercel 서버리스 함수 — 업로드한 문서에서 '문답 텍스트'를 뽑는다.

.docx(Word)·.pdf·.txt/.md·.json 을 받아 편집·재평가용 순수 텍스트로 변환한다.
브라우저가 파일을 base64로 실어 보내면(요청 본문 JSON) 확장자로 형식을 판별한다.
문서 내용은 '저장한 문답'이라는 전제다(원칙: 먼저 문답, 그다음 평가).

요청(JSON): {"filename": "대화.docx", "content_b64": "..."}
응답(JSON): {"text": "추출한 텍스트", "chars": 1234}
"""
import base64
import io
import json
from http.server import BaseHTTPRequestHandler

MAX_BYTES = 8 * 1024 * 1024   # 8MB 상한(디코드 후 기준) — 과도한 문서 방지


def _extract_docx(data):
    from docx import Document                       # python-docx
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # 표 안의 텍스트도 포함(문답을 표로 저장한 경우 대비).
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def _extract_pdf(data):
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n\n".join(pages).strip()


def _extract_json(data):
    """우리 저장 형식(.json)이면 문답 전문(transcript_text)·아이디어를 꺼낸다."""
    obj = json.loads(data.decode("utf-8", "replace"))
    if isinstance(obj, str):
        return obj.strip()
    if isinstance(obj, dict):
        return str(obj.get("transcript_text") or obj.get("idea") or "").strip()
    return ""


def _extract(filename, data):
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    if ext == "docx":
        return _extract_docx(data)
    if ext == "pdf":
        return _extract_pdf(data)
    if ext in ("txt", "md"):
        return data.decode("utf-8", "replace").strip()
    if ext == "json":
        return _extract_json(data)
    raise ValueError(f"지원하지 않는 형식입니다(.{ext}). .json/.docx/.pdf/.txt 만 됩니다.")


class handler(BaseHTTPRequestHandler):
    def _send(self, code, obj):
        body = json.dumps(obj, ensure_ascii=False).encode("utf-8")
        self.send_response(code)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_POST(self):
        try:
            length = int(self.headers.get("content-length", 0) or 0)
            payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}") if length else {}
        except Exception:
            self._send(400, {"error": "요청 본문(JSON)을 읽지 못했습니다."})
            return

        filename = (payload.get("filename") or "").strip()
        b64 = payload.get("content_b64") or ""
        if not filename or not b64:
            self._send(400, {"error": "filename 또는 파일 내용이 비어 있습니다."})
            return
        try:
            data = base64.b64decode(b64)
        except Exception:
            self._send(400, {"error": "파일 디코딩에 실패했습니다."})
            return
        if len(data) > MAX_BYTES:
            self._send(400, {"error": "파일이 너무 큽니다(8MB 이하만 됩니다)."})
            return

        try:
            text = _extract(filename, data)
        except ImportError:
            self._send(500, {"error": "서버에 문서 추출 라이브러리가 설치되지 않았습니다."})
            return
        except ValueError as e:
            self._send(400, {"error": str(e)})
            return
        except Exception as e:
            self._send(422, {"error": f"문서에서 텍스트를 추출하지 못했습니다: {e}"})
            return

        if not text:
            self._send(422, {"error": "문서에서 텍스트를 찾지 못했습니다(스캔 이미지 PDF 등일 수 있습니다)."})
            return
        self._send(200, {"text": text, "chars": len(text)})
