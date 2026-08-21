#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""방법론 .md → .docx 재생성 (기존 문서 서식과 맞춤).

## → Title, ### → Heading 1, #### → Heading 2, 표·목록·인용·코드블록 처리.
줄바꿈으로 끊긴 **굵게**·*기울임*·`코드` 표시가 갈라지지 않도록,
한 블록의 줄을 **먼저 모두 이어붙인 뒤** 런으로 나눈다.
"""
import re
import sys
from docx import Document
from docx.shared import Pt

SRC, DST = sys.argv[1], sys.argv[2]

doc = Document()
n = doc.styles["Normal"]
n.font.name = "맑은 고딕"
n.font.size = Pt(10.5)
STYLES = {x.name for x in doc.styles}

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`.+?`)", re.S)


def runs(p, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            p.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
        else:
            p.add_run(part)


def cells_of(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


lines = open(SRC, encoding="utf-8").read().split("\n")
pending = None            # (style, [줄, ...])


def flush():
    """모아둔 줄을 한 문장으로 이어 붙인 뒤 한 문단으로 쓴다."""
    global pending
    if not pending:
        return
    style, buf = pending
    pending = None
    text = " ".join(x.strip() for x in buf if x.strip())
    if not text:
        return
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    runs(p, text)


def start(style, first):
    global pending
    flush()
    pending = (style, [first])


i = 0
while i < len(lines):
    ln = lines[i]
    s = ln.strip()

    if s.startswith("```"):                                   # 코드블록
        flush()
        i += 1
        buf = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            buf.append(lines[i]); i += 1
        r = doc.add_paragraph().add_run("\n".join(buf))
        r.font.name = "Consolas"; r.font.size = Pt(9)
        i += 1
        continue

    if s.startswith("|") and i + 1 < len(lines) and \
            re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):   # 표
        flush()
        head = cells_of(s)
        i += 2
        body = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            body.append(cells_of(lines[i])); i += 1
        t = doc.add_table(rows=1, cols=len(head))
        t.style = "Light Grid Accent 1"
        for c, h in zip(t.rows[0].cells, head):
            c.text = ""; runs(c.paragraphs[0], h)
        for row in body:
            for c, v in zip(t.add_row().cells, row + [""] * (len(head) - len(row))):
                c.text = ""; runs(c.paragraphs[0], v)
        doc.add_paragraph()
        continue

    if not s:
        flush()
    elif s.startswith("### "):        # 원본 문서 대응: # → Title, ## → H1, ### → H2
        flush(); doc.add_heading(s[4:].strip(), 2)
    elif s.startswith("## "):
        flush(); doc.add_heading(s[3:].strip(), 1)
    elif s.startswith("# "):
        flush(); doc.add_heading(s[2:].strip(), 0)
    elif s.startswith("> "):
        style = "Intense Quote" if "Intense Quote" in STYLES else None
        if pending and pending[0] == style:                    # 여러 줄 인용은 한 문단
            pending[1].append(s[2:])
        else:
            start(style, s[2:])
    elif re.match(r"^[-*] ", s):
        start("List Bullet", s[2:])
    elif re.match(r"^\d+\. ", s):
        start("List Number", re.sub(r"^\d+\. ", "", s))
    elif set(s) <= set("-—=") and len(s) >= 3:
        flush()
    elif pending and ln.startswith("  "):                      # 목록·인용 이어지는 줄
        pending[1].append(s)
    elif pending and pending[0] is None:                       # 본문 이어지는 줄
        pending[1].append(s)
    else:
        start(None, s)
    i += 1

flush()
doc.save(DST)
print(f"저장: {DST}  (문단 {len(doc.paragraphs)} · 표 {len(doc.tables)})")
