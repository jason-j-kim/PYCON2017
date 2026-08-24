#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""서버 설치 안내서를 만든다 — 윈도우 서버에 올리는 쪽.

받는 사람이 다르다. 개인용 안내서는 파이썬을 처음 까는 분을 상대로 하지만,
이쪽은 기관 서버를 맡은 분이다. 윈도우 서버와 IIS 는 아신다고 본다. 그래서
"명령창을 여세요" 는 적지 않는다.

대신 이 시스템에만 있는 것들을 적는다. 왜 personal 이어야 하는지,
왜 --proxy-headers 를 붙이면 안 되는지, /records 가 왜 막혀 있는지,
왜 NSSM 이 아니라 작업 스케줄러인지.
서버를 아는 분일수록 그 옵션을 습관적으로 붙이기 때문이다.

서식은 앞서 만든 문서들과 같다. 색 없음, 음영 없음, 글자 크기 두 가지.
"""
import os
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

KO = "맑은 고딕"
BODY, HEAD, SMALL, MONO = Pt(10.5), Pt(12), Pt(9.5), Pt(9)
BLACK = RGBColor(0, 0, 0)

doc = Document()
sec = doc.sections[0]
sec.top_margin, sec.bottom_margin = Cm(2.4), Cm(2.2)
sec.left_margin, sec.right_margin = Cm(2.6), Cm(2.6)

st = doc.styles["Normal"]
st.font.name = KO
st.font.size = BODY
st.font.color.rgb = BLACK
st.element.rPr.rFonts.set(qn("w:eastAsia"), KO)
st.paragraph_format.line_spacing = Pt(16.5)
st.paragraph_format.space_after = Pt(6.5)


def style_run(r, size=BODY, bold=False):
    r.font.name = KO
    r.font.size = size
    r.font.bold = bold
    r.font.color.rgb = BLACK
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {}); rpr.insert(0, rf)
    for a in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rf.set(qn(a), KO)
    return r


def p(text="", size=BODY, bold=False, after=6.5, before=0, indent=0,
      hang=None, align=None, line=16.5, keep=False):
    par = doc.add_paragraph()
    f = par.paragraph_format
    f.space_after, f.space_before, f.line_spacing = Pt(after), Pt(before), Pt(line)
    if indent:
        f.left_indent = Cm(indent)
    if hang is not None:
        f.first_line_indent = Cm(-hang)
    if align is not None:
        par.alignment = align
    f.keep_together = True
    if keep:
        f.keep_with_next = True
    if text:
        style_run(par.add_run(text), size, bold)
    return par


def rich(chunks, after=6.5, before=0, indent=0, hang=None, keep=False):
    par = p("", after=after, before=before, indent=indent, hang=hang, keep=keep)
    for t, b in chunks:
        style_run(par.add_run(t), BODY, b)
    return par


def h1(text):
    par = p(text, HEAD, True, after=9, before=15, keep=True)
    par.paragraph_format.page_break_before = True
    return par


def h2(text):
    return p(text, BODY, True, after=5, before=13, keep=True)


def bullet(text, indent=0.6):
    return p(f"·  {text}", indent=indent, hang=0.55, after=3.5)


def step(n, text):
    return rich([(f"{n}  ", True), (text, False)], indent=0.75, hang=0.75, after=5)


def rule_line():
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(2)
    par.paragraph_format.space_after = Pt(8)
    pPr = par._p.get_or_add_pPr()
    b = pPr.makeelement(qn("w:pBdr"), {})
    bot = b.makeelement(qn("w:bottom"), {})
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:color"), "000000"); bot.set(qn("w:space"), "1")
    b.append(bot); pPr.append(b)


def table(rows, widths):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    for gc, w in zip(t._tbl.find(qn("w:tblGrid")), widths):
        gc.set(qn("w:w"), str(w))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.width = Pt(widths[ci] / 20)
            for li, line in enumerate(str(val).split("\n")):
                par = cell.paragraphs[0] if li == 0 else cell.add_paragraph()
                par.paragraph_format.space_after = Pt(1)
                par.paragraph_format.space_before = Pt(1)
                par.paragraph_format.line_spacing = Pt(13)
                par.paragraph_format.keep_with_next = (ri < len(rows) - 1)
                style_run(par.add_run(line), SMALL, ri == 0)
        t.rows[ri]._tr.get_or_add_trPr().append(
            t.rows[ri]._tr.makeelement(qn("w:cantSplit"), {}))
    tblPr = t._tbl.tblPr
    b = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge, s in (("top", "4"), ("bottom", "4"), ("left", "2"), ("right", "2"),
                    ("insideH", "2"), ("insideV", "2")):
        e = b.makeelement(qn("w:" + edge), {})
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), s)
        e.set(qn("w:color"), "000000"); e.set(qn("w:space"), "0")
        b.append(e)
    tblPr.append(b)
    t.rows[0]._tr.get_or_add_trPr().append(
        t.rows[0]._tr.makeelement(qn("w:tblHeader"), {}))
    p("", after=10)
    return t


def shell(lines, indent=0.7):
    """명령과 화면 출력. 그대로 옮겨 칠 수 있게 한 덩어리로 둔다."""
    for i, ln in enumerate(lines):
        par = p(ln or " ", MONO, indent=indent, after=0, line=13)
        par.paragraph_format.keep_together = True
        if i < len(lines) - 1:
            par.paragraph_format.keep_with_next = True
    p("", after=5)


C = WD_ALIGN_PARAGRAPH.CENTER

# ══ 표지 ═══════════════════════════════════════════════════════════════
for _ in range(4):
    p("", after=0)
p("정책 아이디어 평가 시스템", SMALL, align=C, after=10)
p("윈도우 서버 설치와 운영", HEAD, True, align=C, after=8)
p("기관 서버에 올려 밖에 공개하는 경우", align=C, after=4)
p("서버를 맡으신 분을 위해", align=C, after=40)
p("서버용 판  ―  방문자가 각자 자기 Claude 키를 넣습니다", align=C, after=6)
for _ in range(6):
    p("", after=0)
p("국민대학교  김재준", HEAD, True, align=C, after=8)
p("joyof15@gmail.com", align=C, after=6)
p("막히시면 위 주소로 연락 주십시오", SMALL, align=C, after=22)
p("2026년 8월판  ·  윈도우 서버용", align=C, after=6)
p("정책아이디어평가_서버용_윈도우_(날짜).zip 과 함께 받으신 문서입니다",
  SMALL, align=C, after=6)
p("이 문서는 zip 안에도 같이 들어 있습니다", SMALL, align=C)

# ══ 0. 착수 전 ═════════════════════════════════════════════════════════
h1("먼저 — 착수 전에 확인할 세 가지")

p("이 문서만 받으셨을 수 있어 필요한 것을 여기 다 적었습니다. 따로 물어보실 "
  "곳 없이 이 문서와 zip 만으로 끝까지 가실 수 있습니다.")

rich([("설치를 시작하기 전에 셋만 확인해 주십시오. ", True),
      ("셋 다 서버 담당자가 아니면 풀 수 없는 것이고, 설치를 다 끝낸 뒤에 "
       "발견하면 원인을 찾기 어렵습니다.", False)])

table([
    ["", "확인할 것", "안 되어 있으면"],
    ["①", "방화벽에서 api.anthropic.com 으로 나가는 통신",
     "아무것도 되지 않습니다.\n망 담당자에게 이 주소를 열어 달라고 하십시오."],
    ["②", "IIS 모듈 두 개\nURL Rewrite · Application Request Routing",
     "앞단 규칙이 조용히 무시되고\n404 만 납니다."],
    ["③", "도메인 하나와 HTTPS 인증서",
     "방문자가 이 화면에 자기 API 키를 붙여넣습니다.\n평문으로 받으면 그 키가 도중에 드러납니다."],
], [500, 3500, 4200])

p("①은 지금 바로 확인하실 수 있습니다. 서버에서 명령창을 열고:")
shell(["curl -sI https://api.anthropic.com/v1/messages -X POST"])
p("401 이 나오면 정상입니다 — 키를 안 보냈으니 거부하는 것이고, 통신 자체는 "
  "닿았다는 뜻입니다. 아무 응답도 없으면 막혀 있는 것입니다.", indent=0.5)

h2("누가 무엇을 하나")

table([
    ["", "하는 일", "필요한 것"],
    ["서버 담당자", "이 문서대로 설치하고 IIS 를 겁니다.\n한 번 해 두면 그다음은 손댈 일이 없습니다.",
     "서버 관리자 권한"],
    ["연구 담당자", "접속 코드를 방문자에게 알립니다.\n쌓인 문답을 내려받아 분석합니다.",
     "서버 원격 접속"],
    ["방문자", "브라우저로 들어와 자기 Claude 키를 넣고\n12문답을 합니다.",
     "Anthropic API 키\n(각자 본인 것)"],
], [1400, 4600, 2200])

rich([("KDI 는 Claude 요금을 내지 않습니다. ", True),
      ("이 판은 방문자가 각자 자기 키를 넣는 방식입니다. 기관 카드로 결제될 "
       "일이 없습니다. 자세한 것은 다음 장에 있습니다.", False)])

# ══ 1 ══════════════════════════════════════════════════════════════════
h1("이 시스템이 무엇인가 — 서버에 올리기 전에")

p("국민이 낸 정책 아이디어를 12번의 문답으로 구조화하고, 국가 데이터 네 "
  "종류로 선례를 확인해 평가하는 웹 시스템입니다. 질문을 만들고 답을 "
  "채점하는 일은 Claude 가 합니다.")

table([
    ["", ""],
    ["구성", "FastAPI + uvicorn. 정적 화면 세 장. 파이썬만 씁니다.\n"
             ".NET·IIS 응용프로그램 풀은 쓰지 않습니다."],
    ["데이터", "폴더 안 SQLite 2개와 JSON 1개, 합쳐서 44MB.\n"
               "외부 DB 를 붙일 필요가 없습니다."],
    ["쌓이는 것", "webapp\\sessions.db (SQLite). 문답·채점·판정.\n"
                  "이것이 연구 자료입니다."],
    ["나가는 통신", "api.anthropic.com (필수)\n"
                    "open.assembly.go.kr (의안 통로를 쓸 때만)"],
    ["동시 접속", "동시에 진행 중인 Claude 호출 40개까지 받습니다\n"
                  "(FastAPI 기본 스레드풀). 한 턴이 10~30초 걸리고\n"
                  "사람은 그사이 생각하므로, 실제 이용자는 그보다\n"
                  "훨씬 많아도 됩니다."],
], [1500, 6700])

h2("리눅스판과 무엇이 다른가")

p("코드와 데이터는 같습니다. 다른 것은 서비스로 만드는 방법과 앞단 "
  "프록시뿐입니다.")

table([
    ["", "리눅스", "윈도우 (이 문서)"],
    ["서비스", "systemd", "작업 스케줄러"],
    ["앞단", "nginx", "IIS (URL Rewrite + ARR)"],
    ["인증서", "certbot", "IIS 관리자에서 바인딩"],
    ["설정 파일", "/etc/policy-eval.env", "keys.local.bat"],
], [1400, 3000, 3800])

h2("판이 둘 있고, 밖에 열면 답은 하나입니다")

table([
    ["", "personal (이 zip)", "experiment"],
    ["방문자가 넣는 것", "초대 코드 + 자기 API 키", "초대 코드만"],
    ["Claude 요금", "방문자 각자", "기관 키 하나에 전부"],
    ["화면", "「Claude 연결」 칸이 보임", "키 칸이 안 보임"],
], [1700, 3300, 3200])

rich([("밖에 여신다면 personal 이어야 합니다. ", True),
      ("초대 코드는 참여자 전원이 아는 값이라 반드시 샙니다 — 메일 한 번 "
       "전달되면 끝입니다. experiment 로 열어 두면 코드가 새는 순간 남의 "
       "실험이 기관 카드로 결제됩니다. 아이디어 1건에 300~600원이므로 "
       "하룻밤이면 감당하기 어려운 금액이 됩니다. 그래서 personal 판에서는 "
       "keys.local.bat 에 ANTHROPIC_API_KEY 를 ", False),
      ("넣지 마십시오", True),
      (" — 넣으면 방문자가 키를 안 넣어도 그 키로 돌아갑니다. "
       "서버_1_설치.bat 이 이 경우를 발견하면 경고합니다.", False)])

# ══ 2 ══════════════════════════════════════════════════════════════════
h1("미리 준비할 것")

table([
    ["", "", ""],
    ["Python", "3.10 이상", "설치할 때 Add python.exe to PATH 체크"],
    ["IIS", "", "역할 추가에서 켭니다"],
    ["IIS 모듈", "URL Rewrite", "없으면 규칙이 무시되고 404 가 납니다"],
    ["", "Application Request Routing", "설치 후 Enable proxy 를 켜야 합니다"],
    ["도메인·인증서", "", "HTTPS 는 선택이 아닙니다"],
    ["방화벽", "443 열기", "밖에서 접속하려면"],
    ["", "api.anthropic.com 나가기", "막히면 아무것도 되지 않습니다"],
    ["", "open.assembly.go.kr 나가기", "의안 통로를 쓸 때만 (선택)"],
], [1400, 2900, 3900])

rich([("바깥으로 나가는 통신을 먼저 확인하십시오. ", True),
      ("기관 프록시 뒤에 있는 서버라면 이것이 가장 흔한 실패 지점입니다. "
       "설치를 다 끝낸 뒤에 발견하면 원인을 찾기 어렵습니다.", False)])

p("서버에서 명령창을 열고 다음을 쳐 보십시오. 401 이 나오면 정상입니다 — "
  "키를 안 보냈으니 거부하는 것이고, 통신 자체는 닿았다는 뜻입니다.")

shell(["curl -sI https://api.anthropic.com/v1/messages -X POST"])

p("curl 이 없으면 브라우저로 https://api.anthropic.com 을 열어 보셔도 "
  "됩니다. 무언가 응답이 오면 통신은 되는 것입니다.")

# ══ 3 ══════════════════════════════════════════════════════════════════
h1("설치 — 배치 두 개")

p("압축을 푼 폴더 안에서 순서대로 누릅니다. C 드라이브 아래의 짧은 경로를 "
  "권합니다. 예: C:\\policy-eval")

h2("1) 서버_1_설치.bat")

p("그냥 두 번 누르면 됩니다. 가상환경을 만들고 파이썬 꾸러미를 깔고 접속 "
  "코드를 정합니다. 2~3분 걸립니다.")

shell([
    "================================================================",
    "  정책 아이디어 평가 - 윈도우 서버 설치",
    "  폴더: C:\\policy-eval",
    "================================================================",
    "",
    "[1/4] 파이썬 확인",
    "      3.12.4  (C:\\Python312\\python.exe)",
    "",
    "[2/4] 가상환경과 파이썬 꾸러미",
    "      만듦 - C:\\policy-eval\\.venv",
    "      완료 - fastapi · uvicorn 확인",
    "",
    "[3/4] 운영 설정",
    "      판: personal",
    "      접속 코드를 새로 만들었습니다: k7m2xq4p",
    "      <- 방문자에게 알려 줄 값입니다. 적어 두세요.",
    "      저장 - C:\\policy-eval\\keys.local.bat",
    "",
    "[4/4] 평가에 쓰는 자료",
    "      (1) 재정          2.8MB  바로 작동",
    "      (2) KDI 연구     26.8MB  바로 작동",
    "      (4) 해외사례     14.8MB  바로 작동",
])

rich([("접속 코드를 적어 두십시오. ", True),
      ("이 줄에 한 번만 찍힙니다. 놓치셨으면 keys.local.bat 의 "
       "SOCRATIC_ACCESS_CODE 를 보시면 됩니다. 바꾸시려면 그 파일을 고치고 "
       "작업을 다시 켜면 됩니다.", False)])

h2("2) 서버_2_서비스등록.bat — 관리자 권한으로")

rich([("이것만은 오른쪽 클릭 → [관리자 권한으로 실행] 이어야 합니다. ", True),
      ("그냥 두 번 누르면 등록에 실패하고 그 이유를 화면에 적어 줍니다.",
       False)])

p("부팅 때 자동으로 켜지도록 작업 스케줄러에 등록합니다.")

shell([
    "  등록했습니다 - 작업 이름 PolicyEval",
    "    · 부팅 30초 뒤 자동으로 켜집니다",
    "    · 죽으면 1분 뒤 다시 켭니다 (최대 999번)",
    "    · 로그온하지 않아도 돕니다 (SYSTEM 계정)",
])

h2("왜 NSSM 이 아니라 작업 스케줄러인가")

p("윈도우에서 파이썬 앱을 서비스로 만드는 길은 셋입니다.")

table([
    ["", ""],
    ["NSSM", "가장 흔하지만 외부 실행 파일을 따로 받아야 합니다.\n"
             "기관 서버에 출처가 분명하지 않은 바이너리를 들이는 것은\n"
             "승인이 어렵습니다."],
    ["pywin32", "파이썬 꾸러미를 하나 더 깔고 서비스 클래스를 써야 합니다."],
    ["작업 스케줄러", "윈도우에 이미 들어 있습니다. 받을 것이 없습니다.\n"
                      "이 판이 쓰는 방법입니다."],
], [1500, 6700])

p("작업 스케줄러도 부팅 시 자동 실행 · 실패 시 재시작 · 로그온 없이 실행을 "
  "모두 합니다. services.msc 목록에 안 보인다는 점만 다릅니다.")

p("손으로 켜고 끄시려면:")
shell(["schtasks /Run /TN PolicyEval", "schtasks /End /TN PolicyEval"])

h2("먼저 손으로 켜 보려면 — 서버_3_직접실행.bat")

p("서비스로 등록하기 전에 한 번 켜 보시는 편이 낫습니다. 창을 닫으면 "
  "꺼지고, 오류가 있으면 그 창에 그대로 나옵니다. 등록한 뒤에는 오류가 "
  "화면에 안 보여 원인을 찾기 어렵습니다.")

p("켜진 뒤 서버의 브라우저에서 다음을 열어 보십시오.")
shell(["http://localhost:8000/policy"])

# ══ 4 ══════════════════════════════════════════════════════════════════
h1("IIS 앞단 걸기")

h2("1) 모듈 두 개")

p("먼저 깔아야 합니다. 없으면 web.config 의 규칙이 조용히 무시되고 404 가 "
  "납니다.")

shell([
    "URL Rewrite",
    "  https://www.iis.net/downloads/microsoft/url-rewrite",
    "Application Request Routing",
    "  https://www.iis.net/downloads/microsoft/application-request-routing",
])

h2("2) ARR 을 켜기 — 빠뜨리기 쉬운 곳")

p("설치만으로는 안 되고 한 번 켜 주어야 합니다.")

shell([
    "IIS 관리자 -> 서버 노드",
    "  -> Application Request Routing Cache",
    "  -> Server Proxy Settings",
    "  -> Enable proxy 체크 -> 적용",
])

h2("3) 사이트를 만들고 web.config 를 놓기")

p("사이트를 하나 만듭니다. 실체 경로는 빈 폴더여도 됩니다 — IIS 가 파일을 "
  "직접 내보내지 않고 전부 파이썬 서버로 넘기기 때문입니다.")

p("그 경로에 deploy_win\\web.config 를 복사합니다. 규칙이 셋 들어 있습니다.")

table([
    ["규칙", "하는 일"],
    ["BlockRecords", "/records 와 /api/records 를 404 로 막습니다"],
    ["ForceHttps", "평문을 https 로 보냅니다 — 주석으로 되어 있습니다"],
    ["ToUvicorn", "나머지를 127.0.0.1:8000 으로 넘깁니다"],
], [1900, 6300])

h2("요청 제한은 기본으로 꺼져 있습니다")

p("문답 한 번은 12번을 주고받으므로 사람이 쓰는 속도는 분당 몇 건을 넘지 "
  "않습니다. 이보다 빠른 것은 사람이 아닙니다. 리눅스판(nginx)에는 분당 30회 "
  "제한이 켜져 있는데, IIS 에서 같은 일을 하려면 역할 기능을 하나 더 깔아야 "
  "해서 기본은 꺼 두었습니다.")

p("켜시려면 먼저 역할 기능을 설치하십시오.")
shell([
    "서버 관리자 -> 역할 및 기능 추가 -> 웹 서버(IIS)",
    "  -> 보안 -> IP 및 도메인 제한",
])

p("그다음 web.config 의 dynamicIpSecurity 주석을 푸십시오. 설치하지 않은 채 "
  "주석을 풀면 사이트 전체가 500 을 냅니다.", indent=0.5)

rich([("nginx 는 429 로 답하지만 IIS 는 403 밖에 고를 수 없습니다. ", True),
      ("로그에서 진짜 거부(접속 코드 오류)와 섞이므로, 403 이 갑자기 늘면 "
       "요청 제한 쪽을 먼저 의심하십시오.", False)], indent=0.5)

h2("4) 인증서")

p("IIS 관리자에서 사이트 바인딩(https, 443)에 붙입니다.")

rule_line()

rich([("인증서를 붙인 뒤에 ForceHttps 의 주석을 푸십시오. ", True),
      ("순서가 거꾸로면 접속이 아예 안 됩니다 — 평문으로 온 요청을 https 로 "
       "보내는데 https 가 아직 없기 때문입니다.", False)])

rich([("IIS 를 쓰지 않으신다면 ", True),
      ("run_server.py 를 HOST=0.0.0.0 으로 띄우고 uvicorn 에 인증서를 직접 "
       "물릴 수도 있습니다. 다만 그러면 BlockRecords 겹막기가 없어지고 앱 "
       "안의 잠금장치 하나만 남습니다. 기관 서버라면 IIS 를 앞에 두는 편이 "
       "낫습니다.", False)])

rich([("HTTPS 는 선택이 아닙니다. ", True),
      ("personal 판에서는 방문자가 화면에 자기 Anthropic API 키를 "
       "붙여넣습니다. 평문으로 받으면 그 키가 도중에 드러납니다. 그 키는 "
       "브라우저와 서버 메모리에만 있고 sessions.db 에도 로그에도 남지 "
       "않는데, 그것이 성립하려면 전송 구간이 막혀 있어야 합니다.", False)])

rule_line()

rich([("run_server.py 에 --proxy-headers 를 붙이지 마십시오. ", True),
      ("서버를 아시는 분일수록 습관적으로 붙이는 옵션이라 따로 적어 "
       "둡니다. 그 옵션은 X-Forwarded-For 를 client.host 로 바꿔치기하는데, "
       "위의 잠금장치가 보는 것이 바로 그 값입니다. 붙이면 밖에서 온 요청이 "
       "내부 주소로 위장해 통과합니다.", False)])

p("함께 드린 run_server.py 에는 그 옵션이 없고, 없는 이유도 파일 안에 "
  "주석으로 적어 두었습니다.")

rule_line()

# ══ 6 ══════════════════════════════════════════════════════════════════
h1("켠 뒤에 확인할 것")

h2("서버_4_상태확인.bat")

p("작업이 등록됐는지만 보지 않고 127.0.0.1:8000 에 실제로 물어봅니다. "
  "작업이 '실행 중'이어도 앱이 죽어 있을 수 있기 때문입니다.")

shell([
    "  상태:              실행 중",
    "  마지막 실행 시간:  2026-08-24 09:00:30",
    "  마지막 결과:       267009",
    "",
    "  http://127.0.0.1:8000/api/config -> 200 정상 응답",
])

p("마지막 결과 267009 는 '아직 돌고 있다'는 뜻입니다. 오류가 아닙니다.")

h2("/api/config 가 무엇을 답해야 하나")

table([
    ["보실 것", "이래야 정상"],
    ["mode", "personal — 방문자가 각자 키를 넣는 판"],
    ["access_required", "true — 초대 코드를 물어봅니다"],
    ["has_server_key", "false — 서버에 기관 키가 없습니다.\n"
     "true 면 방문자가 키를 안 넣어도 기관 키로 돌아갑니다"],
    ["sources 셋", "전부 true — 데이터가 제자리에 있습니다"],
], [2100, 6100])

h2("사람 손으로 한 번")

p("브라우저에서 도메인을 열고 접속 코드를 넣은 뒤, 짧은 아이디어로 문답을 "
  "한 번 끝까지 돌려 보십시오. Claude 키가 필요하므로 시험용으로 하나 "
  "준비하셔야 합니다. 문답이 끝나고 점수와 커버리지가 나오면 네 통로가 "
  "모두 살아 있는 것입니다.")

p("12문답에 20~30분, 그 뒤 선례 조사에 1~2분 걸립니다.")

# ══ 7 ══════════════════════════════════════════════════════════════════
h1("운영")

h2("설정을 바꾸려면")

p("전부 keys.local.bat 한 곳에 있습니다. 메모장으로 고친 뒤 작업을 다시 "
  "켜면 됩니다.")

shell(["schtasks /End /TN PolicyEval", "schtasks /Run /TN PolicyEval"])

table([
    ["", "무엇"],
    ["SOCRATIC_ACCESS_CODE", "접속 코드. 방문자에게 알려 줄 값."],
    ["SOCRATIC_MAX_SESSIONS_PER_DAY",
     "하루 세션 상한. 기본 30, 설치가 100 으로 둡니다.\n"
     "밖에 열어 두는 동안의 안전판이므로 필요 이상으로\n"
     "올리지 않는 편이 낫습니다."],
    ["ASSEMBLY_KEY",
     "국회 의안 인증키 (선택). 넣어 두면 방문자가 안 넣어도\n"
     "(3) 통로가 돕니다. 열린국회정보에서 발급합니다 —\n"
     "공공데이터포털(data.go.kr) 키와 다릅니다."],
    ["ANTHROPIC_API_KEY",
     "experiment 판일 때만 채웁니다. personal 판에서 채우면\n"
     "요금이 기관에 붙습니다. 비워 두십시오."],
], [2900, 5300])

rich([("keys.local.bat 에는 키가 들어갑니다. ", True),
      ("이 폴더를 통째로 남에게 주지 마십시오.", False)])

h2("연구 자료")

rich([("webapp\\sessions.db 가 연구 자료입니다. ", True),
      ("문답·채점·선례 판정이 전부 여기 쌓입니다. 백업 대상에 넣으시되, "
       "개인정보가 담길 수 있으므로 배포용 zip 에는 넣지 마십시오.", False)])

p("사람이 읽을 형태로 보려면:")
shell([".venv\\Scripts\\python webapp\\show_session.py"])

h2("갱신")

p("새 zip 을 풀고 webapp\\sessions.db 와 keys.local.bat 을 남겨 둔 채 "
  "나머지를 덮어쓰십시오. 그다음:")
shell([
    "서버_1_설치.bat",
    "schtasks /End /TN PolicyEval",
    "schtasks /Run /TN PolicyEval",
])

rich([("인증을 더 두껍게 하려면 — ", True),
      ("지금은 초대 코드 하나뿐이고 참여자 전원이 아는 값입니다. 소속 "
       "연구원만 쓰게 하시려면 IIS 쪽에 기본 인증이나 기관 SSO 를 한 겹 더 "
       "얹으십시오.", False)], before=10)

# ══ 8 ══════════════════════════════════════════════════════════════════
h1("잘 안 될 때")

table([
    ["이런 일이 생기면", "이렇게 하십시오"],
    ["배치를 눌렀는데 창이\n순식간에 사라진다",
     "파이썬이 없거나 PATH 에 없습니다.\n"
     "python.org 에서 다시 설치하며\n"
     "「Add python.exe to PATH」 를 체크하십시오."],
    ["등록에 실패한다",
     "관리자 권한으로 실행하지 않았습니다.\n"
     "서버_2_서비스등록.bat 오른쪽 클릭 →\n"
     "[관리자 권한으로 실행]"],
    ["작업은 실행 중인데\n화면이 안 뜬다",
     "서버_3_직접실행.bat 으로 켜 보십시오.\n"
     "오류가 화면에 그대로 나옵니다."],
    ["IIS 가 404 를 낸다",
     "URL Rewrite / ARR 모듈이 없거나,\n"
     "ARR 의 Enable proxy 를 안 켠 것입니다."],
    ["502 · 504",
     "파이썬 서버가 죽었습니다.\n"
     "서버_4_상태확인.bat 부터 보십시오."],
    ["화면은 뜨는데 문답이\n시작되지 않는다",
     "방문자가 API 키를 안 넣었거나 키가 거부됐습니다.\n"
     "화면 「Claude 연결」 칸에 이유가 찍힙니다."],
    ["「닿지 못했습니다」",
     "방화벽이 api.anthropic.com 을 막고 있습니다.\n"
     "망 담당자에게 이 주소를 확인하십시오."],
    ["의안 통로가 의심스럽다",
     ".venv\\Scripts\\python webapp\\check_bill.py\n"
     "키 없음 / 망 차단 / 키 거부 / 정상을 구분합니다."],
    ["Claude 키가 의심스럽다",
     ".venv\\Scripts\\python webapp\\check_claude.py\n"
     "붙여넣다 섞여든 공백·보이지 않는 문자를\n"
     "몇 번째 자리인지까지 찍어 줍니다."],
], [2400, 5800])

h2("401 과 400 은 다릅니다")

p("인증 실패(401)는 키 자체의 문제입니다. 잔액 부족과 한도 초과는 401 이 "
  "아니라 400 으로 나옵니다. 401 이면 돈 문제가 아닙니다.")

# ══ 9 ══════════════════════════════════════════════════════════════════
h1("폴더 안에 무엇이 있나")

table([
    ["", ""],
    ["mode.txt", "personal — 이 한 줄이 판을 정합니다"],
    ["서버_1_설치.bat", "가상환경 · 꾸러미 · 접속 코드"],
    ["서버_2_서비스등록.bat", "작업 스케줄러 등록 (관리자 권한)"],
    ["서버_3_직접실행.bat", "손으로 켜 보기. 창을 닫으면 꺼집니다"],
    ["서버_4_상태확인.bat", "떠 있는지 · 실제로 응답하는지"],
    ["deploy_win\\web.config", "IIS 앞단 설정"],
    ["deploy_win\\README.md", "더 자세한 설명"],
    ["data\\fiscal.json", "세출예산 14,122개 사업        2.8MB"],
    ["data\\kdi.sqlite", "KDI 발간물 7,362건          26.8MB"],
    ["data\\opsi_policies.db", "OECD OPSI 1,015건·98개국   14.8MB"],
    ["webapp\\", "웹 서버와 화면. sessions.db 도 여기 생깁니다"],
    ["socratic\\", "문답·채점 엔진과 프롬프트"],
    ["kdinov\\", "KDI 코퍼스 정밀 판정기"],
    ["읽어보세요.txt", "이 문서의 요약본"],
], [2600, 5600])

p("배치 파일은 \"python OO.py\" 한 줄만 합니다. 한국어 안내와 로직은 전부 "
  ".py 안에 있습니다 — 배치에 한국어를 넣으면 윈도우 인코딩에 따라 깨지기 "
  "때문입니다.")

rich([("압축을 풀면 45MB 를 차지합니다. ", True),
      ("zip 은 8.8MB 이지만 데이터가 압축돼 있어서 그렇습니다. 푸는 데 "
       "30초에서 1분쯤 걸립니다.", False)])

rich([("이 문서에 적히지 않은 별도 지시는 없습니다. ", True),
      ("메일이나 구두로 따로 전달받으실 것이 없도록, 당부할 것을 모두 이 "
       "문서 안에 넣었습니다. 그래도 막히시면 국민대학교 김재준 "
       "(joyof15@gmail.com) 으로 연락 주십시오.", False)], before=12)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "서버설치_안내_윈도우.docx")
doc.save(out)
print("저장 완료 —", out)