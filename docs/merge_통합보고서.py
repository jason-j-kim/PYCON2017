#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""세 문서를 하나로 합친다.  온국민_v3 → 정책 1 → 정책 4

그냥 이어 붙이면 앞 23쪽은 흑백인데 뒤부터 색과 음영이 나온다. 게다가
뒤 두 문서는 위계를 색으로만 표현하고 있어(초록=충족, 빨강=미충족,
흰 글씨+음영=근거 배지), 색을 없애면 구조가 통째로 뭉개진다.

그래서 색이 하던 일을 굵기·크기·괄호로 옮긴다.

  문서 제목        빨강 12pt      →  검정 12pt 굵게
  네 아이디어 머리  빨강 10/12/14  →  검정 12pt 굵게 (원본은 크기가 들쭉날쭉)
  판정 밴드        파랑 13.5pt    →  검정 12pt 굵게
  충족 · 미충족    초록 · 빨강    →  검정 굵게 (✓·✗ 기호가 이미 구분한다)
  근거 배지        흰 글씨+음영    →  [A급][재정] 처럼 대괄호로
  옅은 회색 주석    회색           →  검정 (크기로 구분)

정책 4의 요약표는 그림 파일로 들어 있었다. 진짜 표로 다시 만든다 —
검색되고 인쇄가 깨끗하며 나머지 표들과 서식이 같아진다.
"""
import copy
import re
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

KO = "맑은 고딕"
BODY, HEAD, SMALL = Pt(10.5), Pt(12), Pt(10)
LINE = 320
BLACK = RGBColor(0, 0, 0)

base = Document("A_온국민v3.docx")


def style_run(r, size=BODY, bold=False):
    r.font.name = KO
    r.font.size = size
    r.font.bold = bold
    r.font.color.rgb = BLACK
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {}); rpr.insert(0, rf)
    rf.set(qn("w:eastAsia"), KO)
    rf.set(qn("w:ascii"), KO)
    rf.set(qn("w:hAnsi"), KO)
    return r


def new_para(text="", size=BODY, bold=False, after=130, line=LINE, align=None):
    p = base.add_paragraph()
    p.paragraph_format.space_after = Pt(after / 20)
    p.paragraph_format.line_spacing = Pt(line / 20)
    if align is not None:
        p.alignment = align
    if text:
        style_run(p.add_run(text), size, bold)
    return p


def page_break():
    p = base.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ── 1. 표지에 저자 표기 ────────────────────────────────────────────────
cover = base.paragraphs
idx = next(i for i, p in enumerate(cover) if p.text.strip() == "2026")
anchor = cover[idx]
for txt, sz, bold in [("국민대학교  김재준", HEAD, True)]:
    np = copy.deepcopy(anchor._p)
    anchor._p.addprevious(np)
    from docx.text.paragraph import Paragraph
    par = Paragraph(np, anchor._parent)
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    par.paragraph_format.space_after = Pt(20)
    style_run(par.add_run(txt), sz, bold)
# 저자 줄과 연도 사이 간격
np2 = copy.deepcopy(anchor._p)
anchor._p.addprevious(np2)
from docx.text.paragraph import Paragraph as _P
sp = _P(np2, anchor._parent)
for r in list(sp.runs):
    r._element.getparent().remove(r._element)
sp.paragraph_format.space_after = Pt(28)

print("표지에 저자 표기 추가")


# ── 2. 뒤 문서를 정규화해 붙인다 ───────────────────────────────────────
TURN = re.compile(r"^(턴\s*\d+\s*·\s*[^\s(]+(?:\s*\([^)]*\))?)\s*(.*)$", re.S)


def para_kind(p):
    """색·크기로 표현된 위계를 읽어 종류를 돌려준다."""
    r = next((r for r in p.runs if r.text.strip()), None)
    if r is None:
        return "empty", None
    col = str(r.font.color.rgb) if (r.font.color and r.font.color.rgb) else ""
    sz = r.font.size.pt if r.font.size else 10.0
    bold = bool(r.bold)
    if col == "EE0000" or (sz >= 12 and bold and not col):
        return "title", None
    if sz >= 13:
        return "band", None
    if col in ("23639C", "26241F") and bold:
        return "head", None
    if bold:
        return "sub", None
    return "body", None


def append_doc(path, label):
    src = Document(path)
    n_img = 0
    for p in src.paragraphs:
        # 그림은 따로 처리한다(정책 4의 요약표)
        if "blip" in p._p.xml:
            n_img += 1
            add_summary_table()
            continue
        text = p.text.strip()
        if not text:
            new_para("", after=60)
            continue
        kind, _ = para_kind(p)

        if kind == "title":
            new_para(text, HEAD, True, after=200)
            continue
        if kind == "band":
            new_para(text, HEAD, True, after=110)
            continue
        if kind == "head":
            new_para(text, BODY, True, after=110)
            continue

        # 근거 배지(흰 글씨 + 음영)를 대괄호로 옮긴다
        parts = []
        for r in p.runs:
            if not r.text:
                continue
            col = str(r.font.color.rgb) if (r.font.color and r.font.color.rgb) else ""
            parts.append(("badge" if col == "FFFFFF" else "plain", r.text))
        merged, buf, mode = [], "", None
        for m, t in parts:
            if m != mode and buf:
                merged.append((mode, buf)); buf = ""
            mode, buf = m, buf + t
        if buf:
            merged.append((mode, buf))

        if any(m == "badge" for m, _ in merged):
            para = new_para(after=90)
            for m, t in merged:
                style_run(para.add_run(f"[{t.strip()}] " if m == "badge" else t),
                          SMALL, m == "badge")
            continue

        # 문답 전문이면 "턴 N · 화자" 를 굵은 줄로 떼어 낸다
        mt = TURN.match(text)
        if mt and label == "B":
            lab, rest = mt.group(1).strip(), mt.group(2).strip()
            m2 = re.match(r"^(\([^)]*\))\s*(.*)$", rest, re.S)
            if m2 and "(" not in lab:
                lab, rest = lab + " " + m2.group(1), m2.group(2).strip()
            new_para(lab, BODY, True, after=50)
            if rest:
                new_para(rest, BODY, False, after=140)
            continue

        new_para(text, BODY, kind == "sub", after=110 if kind == "sub" else 130)
    return n_img


def add_summary_table():
    """정책 4의 요약표 — 그림으로 들어 있던 것을 진짜 표로 다시 만든다."""
    rows = [
        ["", "독창성(방어)\n35%", "실용성\n35%", "수용태도\n30%", "방어력\n총점", "창의성(축 B)"],
        ["① 마을 데이터 금고", "6.0", "3.0", "4.0", "4.35", "계열 밖 시도 (확신 중)"],
        ["② 돌봄 시간 이전제", "9.0", "7.0", "7.5", "7.85", "계열 밖 시도 (확신 상)"],
        ["③ 청년 창업 플랫폼", "2.5", "3.0", "2.0", "2.52", "선례 명확 (확신 상)"],
        ["④ 배달 수수료 환급", "4.0", "9.0", "9.0", "7.25", "선례 명확 (확신 상)"],
    ]
    widths = [2450, 1150, 950, 1050, 950, 2350]
    t = base.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    # 워드는 셀 폭을 보지만 리브레오피스는 tblGrid 를 본다. 둘 다 맞춘다.
    for gc, w in zip(t._tbl.find(qn("w:tblGrid")), widths):
        gc.set(qn("w:w"), str(w))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.width = Pt(widths[ci] / 20)
            cell.text = ""
            for li, line in enumerate(str(val).split("\n")):
                par = cell.paragraphs[0] if li == 0 else cell.add_paragraph()
                par.paragraph_format.space_after = Pt(0)
                par.paragraph_format.line_spacing = Pt(13)
                if ci and ri:
                    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci < 5 else None
                style_run(par.add_run(line), SMALL, ri == 0 or ci == 0)
    # 표 테두리 — 본문 표와 같은 모양(검정 실선)
    tblPr = t._tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge, sz in (("top", "4"), ("bottom", "4"), ("left", "2"), ("right", "2"),
                     ("insideH", "2"), ("insideV", "2")):
        e = borders.makeelement(qn("w:" + edge), {})
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), sz)
        e.set(qn("w:color"), "000000"); e.set(qn("w:space"), "0")
        borders.append(e)
    tblPr.append(borders)
    new_para("", after=140)
    print("  정책 4의 요약표 그림을 표로 재구성")


page_break()
n1 = append_doc("B_정책1.docx", "B")
page_break()
n2 = append_doc("C_정책4.docx", "C")
print(f"정책 1 붙임 · 정책 4 붙임 (그림 {n1 + n2}개 처리)")

base.save("온국민_아이디어등록제_통합보고서.docx")
print("저장 완료")
