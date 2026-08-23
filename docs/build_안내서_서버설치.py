#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""서버 설치 안내서를 만든다 — 리눅스 서버에 올리는 쪽.

받는 사람이 다르다. 개인용 안내서는 파이썬을 처음 까는 분을 상대로 하지만,
이쪽은 연구소 서버를 맡은 분이다. 리눅스와 nginx 는 아신다고 본다. 그래서
"명령창을 여세요" 는 적지 않는다.

대신 이 시스템에만 있는 것들을 적는다. 왜 personal 이어야 하는지,
왜 --proxy-headers 를 붙이면 안 되는지, /records 가 왜 막혀 있는지.
서버를 아는 분일수록 그 옵션을 습관적으로 붙이기 때문이다.

서식은 앞서 만든 문서들과 같다. 색 없음, 음영 없음, 글자 크기 두 가지.
"""
import os
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

KO = "맑은 고딕"
BODY, HEAD, SMALL, MONO = Pt(10.5), Pt(12), Pt(9.5), Pt(9)
BLACK = RGBColor(0, 0, 0)

doc = Document()
sec = doc.sections[0]
sec.top_margin, sec.bottom_margin = Cm(2.4), Cm(2.2)
sec.left_margin, sec.right_margin = Cm(2.6), Cm(2.6)

st = doc.styles["Normal"]
st.font.name = KO
st.font.size = BODY
st.font.color.rgb = BLACK
st.element.rPr.rFonts.set(qn("w:eastAsia"), KO)
st.paragraph_format.line_spacing = Pt(16.5)
st.paragraph_format.space_after = Pt(6.5)


def style_run(r, size=BODY, bold=False):
    r.font.name = KO
    r.font.size = size
    r.font.bold = bold
    r.font.color.rgb = BLACK
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {}); rpr.insert(0, rf)
    for a in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rf.set(qn(a), KO)
    return r


def p(text="", size=BODY, bold=False, after=6.5, before=0, indent=0,
      hang=None, align=None, line=16.5, keep=False):
    par = doc.add_paragraph()
    f = par.paragraph_format
    f.space_after, f.space_before, f.line_spacing = Pt(after), Pt(before), Pt(line)
    if indent:
        f.left_indent = Cm(indent)
    if hang is not None:
        f.first_line_indent = Cm(-hang)
    if align is not None:
        par.alignment = align
    f.keep_together = True
    if keep:
        f.keep_with_next = True
    if text:
        style_run(par.add_run(text), size, bold)
    return par


def rich(chunks, after=6.5, before=0, indent=0, hang=None, keep=False):
    par = p("", after=after, before=before, indent=indent, hang=hang, keep=keep)
    for t, b in chunks:
        style_run(par.add_run(t), BODY, b)
    return par


def h1(text):
    par = p(text, HEAD, True, after=9, before=15, keep=True)
    par.paragraph_format.page_break_before = True
    return par


def h2(text):
    return p(text, BODY, True, after=5, before=13, keep=True)


def bullet(text, indent=0.6):
    return p(f"·  {text}", indent=indent, hang=0.55, after=3.5)


def step(n, text):
    return rich([(f"{n}  ", True), (text, False)], indent=0.75, hang=0.75, after=5)


def rule_line():
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(2)
    par.paragraph_format.space_after = Pt(8)
    pPr = par._p.get_or_add_pPr()
    b = pPr.makeelement(qn("w:pBdr"), {})
    bot = b.makeelement(qn("w:bottom"), {})
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:color"), "000000"); bot.set(qn("w:space"), "1")
    b.append(bot); pPr.append(b)


def table(rows, widths):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    for gc, w in zip(t._tbl.find(qn("w:tblGrid")), widths):
        gc.set(qn("w:w"), str(w))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.width = Pt(widths[ci] / 20)
            for li, line in enumerate(str(val).split("\n")):
                par = cell.paragraphs[0] if li == 0 else cell.add_paragraph()
                par.paragraph_format.space_after = Pt(1)
                par.paragraph_format.space_before = Pt(1)
                par.paragraph_format.line_spacing = Pt(13)
                par.paragraph_format.keep_with_next = (ri < len(rows) - 1)
                style_run(par.add_run(line), SMALL, ri == 0)
        t.rows[ri]._tr.get_or_add_trPr().append(
            t.rows[ri]._tr.makeelement(qn("w:cantSplit"), {}))
    tblPr = t._tbl.tblPr
    b = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge, s in (("top", "4"), ("bottom", "4"), ("left", "2"), ("right", "2"),
                    ("insideH", "2"), ("insideV", "2")):
        e = b.makeelement(qn("w:" + edge), {})
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), s)
        e.set(qn("w:color"), "000000"); e.set(qn("w:space"), "0")
        b.append(e)
    tblPr.append(b)
    t.rows[0]._tr.get_or_add_trPr().append(
        t.rows[0]._tr.makeelement(qn("w:tblHeader"), {}))
    p("", after=10)
    return t


def shell(lines, indent=0.7):
    """명령과 화면 출력. 그대로 옮겨 칠 수 있게 한 덩어리로 둔다."""
    for i, ln in enumerate(lines):
        par = p(ln or " ", MONO, indent=indent, after=0, line=13)
        par.paragraph_format.keep_together = True
        if i < len(lines) - 1:
            par.paragraph_format.keep_with_next = True
    p("", after=5)


C = WD_ALIGN_PARAGRAPH.CENTER

# ══ 표지 ═══════════════════════════════════════════════════════════════
for _ in range(4):
    p("", after=0)
p("정책 아이디어 평가 시스템", SMALL, align=C, after=10)
p("서버 설치와 운영", HEAD, True, align=C, after=8)
p("리눅스 서버에 올려 밖에 공개하는 경우", align=C, after=4)
p("서버를 맡으신 분을 위해", align=C, after=40)
p("서버용 판  ―  방문자가 각자 자기 Claude 키를 넣습니다", align=C, after=6)
for _ in range(6):
    p("", after=0)
p("국민대학교  김재준", HEAD, True, align=C, after=24)
p("2026", align=C)

# ══ 1. 무엇을 올리나 ═══════════════════════════════════════════════════
h1("이 시스템이 무엇인가 — 서버에 올리기 전에")

p("국민이 낸 정책 아이디어를 12번의 문답으로 구조화하고, 국가 데이터 네 "
  "종류로 선례를 확인해 평가하는 웹 시스템입니다. 질문을 만들고 답을 "
  "채점하는 일은 Claude 가 합니다.")

p("서버 쪽에서 알아 두실 것은 다음과 같습니다.")

table([
    ["", ""],
    ["구성", "FastAPI + uvicorn. 정적 화면 세 장. 파이썬만 씁니다."],
    ["데이터", "폴더 안 SQLite 2개와 JSON 1개, 합쳐서 44MB.\n"
               "외부 DB 를 붙일 필요가 없습니다."],
    ["쌓이는 것", "webapp/sessions.db (SQLite). 문답·채점·판정.\n"
                  "이것이 연구 자료입니다."],
    ["나가는 통신", "api.anthropic.com (필수)\n"
                    "open.assembly.go.kr (의안 통로를 쓸 때만)"],
    ["동시 접속", "동시에 진행 중인 Claude 호출 40개까지 받습니다\n"
                  "(FastAPI 기본 스레드풀). 한 턴이 10~30초 걸리고\n"
                  "사람은 그사이 생각하므로, 실제 이용자는 그보다\n"
                  "훨씬 많아도 됩니다."],
], [1500, 6700])

h2("판이 둘 있고, 밖에 열면 답은 하나입니다")

p("mode.txt 한 줄이 판을 가릅니다. 갈리는 것은 누가 Claude 요금을 "
  "내는가입니다.")

table([
    ["", "personal (이 zip)", "experiment"],
    ["방문자가 넣는 것", "초대 코드 + 자기 API 키", "초대 코드만"],
    ["Claude 요금", "방문자 각자", "연구소 키 하나에 전부"],
    ["화면", "「Claude 연결」 칸이 보임", "키 칸이 안 보임"],
], [1700, 3300, 3200])

rich([("밖에 여신다면 personal 이어야 합니다. ", True),
      ("초대 코드는 참여자 전원이 아는 값이라 반드시 샙니다 — 메일 한 번 "
       "전달되면 끝입니다. experiment 로 열어 두면 코드가 새는 순간 남의 "
       "실험이 연구소 카드로 결제됩니다. 아이디어 1건에 300~600원이므로 "
       "하룻밤이면 감당하기 어려운 금액이 됩니다.", False)])

p("experiment 가 맞는 경우는 초대한 사람만 들어오는 닫힌 실험입니다. "
  "그때도 콘솔에서 그 키에 월 사용 한도를 걸어야 합니다. 이 zip 은 "
  "personal 로 되어 있습니다.")

# ══ 2. 준비 ════════════════════════════════════════════════════════════
h1("미리 준비할 것")

table([
    ["", "", ""],
    ["Python", "3.10 이상", "대부분의 배포판에 이미 있습니다"],
    ["nginx", "", "앞단 역방향 프록시"],
    ["certbot", "", "무료 인증서. HTTPS 는 선택이 아닙니다"],
    ["도메인", "예: policy.example.ac.kr", "인증서를 받으려면 필요합니다"],
    ["방화벽", "443 열기", "밖에서 접속하려면"],
    ["", "api.anthropic.com 나가기", "막히면 아무것도 되지 않습니다"],
    ["", "open.assembly.go.kr 나가기", "의안 통로를 쓸 때만 (선택)"],
], [1300, 3000, 3900])

rich([("바깥으로 나가는 통신을 먼저 확인하십시오. ", True),
      ("폐쇄망이나 프록시 뒤에 있는 서버라면 이것이 가장 흔한 실패 "
       "지점입니다. 설치를 다 끝낸 뒤에 발견하면 원인을 찾기 어렵습니다.",
       False)])

p("서버에서 다음이 통하는지 먼저 보십시오. 401 이 나오면 정상입니다 — "
  "키를 안 보냈으니 거부하는 것이고, 통신 자체는 닿았다는 뜻입니다.")

shell(["curl -sI https://api.anthropic.com/v1/messages -X POST"])

# ══ 3. 설치 ════════════════════════════════════════════════════════════
h1("설치 — 두 줄")

p("압축을 푼 폴더로 들어가 스크립트를 돌립니다. /opt/policy-eval 로 "
  "옮기는 것까지 스크립트가 합니다.")

shell([
    "unzip 정책아이디어평가_서버용_*.zip",
    "cd 정책아이디어평가",
    "sudo bash deploy/setup_server.sh",
])

p("스크립트가 하는 일은 여섯 가지입니다. 이미 되어 있는 것은 건너뛰므로 "
  "여러 번 돌려도 안전합니다.")

table([
    ["", "무엇을"],
    ["1", "전용 계정 policyeval 을 만듭니다. root 로 돌리지 않기 위해서입니다."],
    ["2", "가상환경을 만들고 fastapi · uvicorn 을 깝니다."],
    ["3", "mode.txt 를 정합니다 (기본 personal)."],
    ["4", "/etc/policy-eval.env 를 만들고 600 으로 잠급니다.\n"
          "접속 코드를 무작위로 만들어 화면에 찍어 줍니다."],
    ["5", "systemd 유닛을 걸고 켭니다."],
    ["6", "남은 것(nginx · 인증서)을 명령까지 찍어 알려 줍니다."],
], [700, 7500])

rich([("sessions.db 는 덮어쓰지 않습니다. ", True),
      ("갱신하실 때 이미 쌓인 연구 자료가 날아가지 않도록 rsync 에서 "
       "제외해 두었습니다. /etc/policy-eval.env 도 이미 있으면 그대로 "
       "둡니다.", False)])

h2("설치가 끝나면 찍히는 것")

shell([
    "[1/6] 계정 policyeval 만듦",
    "[2/6] 코드 배치와 파이썬 꾸러미",
    "      파이썬 3.11",
    "      fastapi · uvicorn 준비됨",
    "[3/6] 판: personal",
    "[4/6] /etc/policy-eval.env 만듦 (600)",
    "      접속 코드: k7m2xq4p   ← 방문자에게 알려 줄 값",
    "[5/6] policy-eval 서비스 실행 중",
])

rich([("접속 코드를 적어 두십시오. ", True),
      ("이 줄에 한 번만 찍힙니다. 놓치셨으면 /etc/policy-eval.env 의 "
       "SOCRATIC_ACCESS_CODE 를 보시면 됩니다. 바꾸시려면 그 파일을 고치고 "
       "서비스를 다시 켜면 됩니다.", False)])

# ══ 4. nginx ═══════════════════════════════════════════════════════════
h1("nginx 걸기")

step("1", "설정을 옮기고 도메인을 바꿉니다.")
shell([
    "sudo cp /opt/policy-eval/deploy/nginx.conf \\",
    "        /etc/nginx/sites-available/policy-eval",
    "sudo vi /etc/nginx/sites-available/policy-eval",
    "#   server_name 을 실제 도메인으로",
])

step("2", "걸고 확인합니다.")
shell([
    "sudo ln -s /etc/nginx/sites-available/policy-eval \\",
    "           /etc/nginx/sites-enabled/",
    "sudo nginx -t && sudo systemctl reload nginx",
])

step("3", "인증서를 받습니다.")
shell(["sudo certbot --nginx -d policy.example.ac.kr"])

rich([("함께 드린 nginx.conf 에는 443 블록이 없습니다. ", True),
      ("일부러 그렇게 두었습니다. 인증서를 아직 받지 않았는데 "
       "listen 443 ssl 을 써 두면 nginx 가 이렇게 답하며 뜨지 "
       "않습니다.", False)])

shell(['no "ssl_certificate" is defined for the "listen ... ssl" directive'])

p("그러면 certbot 도 돌릴 수 없습니다 — certbot --nginx 는 nginx 를 다시 "
  "읽어야 하는데 설정이 깨져 있으면 거기서 멈춥니다. 닭과 달걀이 됩니다.")

p("그래서 80 만 두고 시작합니다. 위 3번의 certbot 이 이 블록을 제자리에서 "
  "고쳐 443·인증서·평문 이동을 알아서 붙입니다. 아래에 적어 둔 location "
  "들은 그대로 옮겨갑니다. 손으로 고치실 것은 없습니다.")

rule_line()

rich([("HTTPS 는 선택이 아닙니다. ", True),
      ("personal 판에서는 방문자가 화면에 자기 Anthropic API 키를 "
       "붙여넣습니다. 평문(http)으로 받으면 그 키가 도중에 드러납니다. "
       "인증서 없이 열지 마십시오.", False)])

p("방문자의 키는 그 브라우저와 서버 메모리에만 있습니다. sessions.db 에도 "
  "로그에도 남지 않습니다. 이것이 성립하려면 전송 구간이 막혀 있어야 "
  "합니다.")

rule_line()

h2("nginx 설정에 들어 있는 것")

table([
    ["", "왜"],
    ["80 → 443 강제 이동", "certbot 이 붙입니다. 평문으로 들어온 것을\n그냥 받지 않습니다."],
    ["/records · /api/records\n→ 404", "기록 화면을 밖에서 막습니다.\n"
     "다음 장에 이유를 적었습니다."],
    ["분당 30회 제한", "문답 한 번은 12번을 주고받으므로 사람이 쓰는\n"
     "속도는 이를 넘지 않습니다."],
    ["읽기 타임아웃 300초", "Claude 호출이 한 턴에 30초를 넘길 때가\n"
     "있습니다. 기본 60초로는 끊깁니다."],
    ["proxy_buffering off", "답이 오는 대로 흘려보냅니다."],
], [2300, 5900])

# h2 + 문단으로 두면 "다음과 붙여 둠" 이 걸려 통째로 다음 쪽으로 넘어간다.
# 장 끝의 짧은 덧말이므로 한 문단으로 둔다.
rich([("인증을 더 두껍게 하려면 — ", True),
      ("지금은 초대 코드 하나뿐이고 참여자 전원이 아는 값입니다. 소속 "
       "연구원만 쓰게 하시려면 nginx 쪽에 한 겹 더 얹으십시오. 기본 "
       "인증(htpasswd)이 간단합니다.", False)], before=10)


# ══ 5. 기록 화면 ═══════════════════════════════════════════════════════
h1("기록 화면은 밖에서 열리지 않습니다")

p("이것은 고장이 아니라 설계입니다. sessions.db 에는 여러 사람의 문답이 "
  "함께 쌓입니다. 접속한 사람이 남의 대화를 받아 갈 수 있으면 안 됩니다. "
  "초대 코드로 막는 것으로는 부족합니다 — 참여자 전원이 아는 값입니다.")

p("두 겹으로 막혀 있습니다.")

table([
    ["어디서", "무엇을 보고", "결과"],
    ["앱 (_require_local)",
     "프록시가 붙인 머리표\nX-Forwarded-For · X-Real-IP\ncf-connecting-ip 등",
     "403"],
    ["nginx", "경로 /records · /api/records", "404"],
], [2200, 4000, 2000])

p("앞의 것이 주소만 보지 않고 머리표를 보는 이유가 있습니다. cloudflared "
  "터널을 쓰면 밖에서 들어온 사람도 서버 입장에서는 127.0.0.1 로 "
  "보입니다. 주소만 보는 잠금장치는 그대로 뚫립니다.")

h2("운영자는 SSH 터널로 봅니다")

shell([
    "ssh -L 8000:127.0.0.1:8000 사용자@서버",
    "#   그다음 내 PC 브라우저에서",
    "#   http://localhost:8000/records",
])

p("프록시를 거치지 않으므로 그대로 통과합니다. 문답 자료가 밖으로 나갈 "
  "길이 아예 없다는 점에서 가장 안전한 방법입니다.")

p("건마다 보고서(.docx) · 문답 전문(.txt) · 원자료(.json) 를 내려받을 수 "
  "있고, 「전부 내려받기」 로 한 번에 받을 수도 있습니다.")

rule_line()

rich([("uvicorn 에 --proxy-headers 를 붙이지 마십시오. ", True),
      ("서버를 아시는 분일수록 습관적으로 붙이는 옵션이라 따로 적어 "
       "둡니다. 그 옵션은 X-Forwarded-For 를 client.host 로 "
       "바꿔치기하는데, 위의 잠금장치가 보는 것이 바로 그 값입니다. "
       "붙이면 밖에서 온 요청이 내부 주소로 위장해 통과합니다.", False)])

p("함께 드린 systemd 유닛에는 그 옵션이 없고, 없는 이유도 파일 안에 "
  "주석으로 적어 두었습니다.")

rule_line()

# ══ 6. 확인 ════════════════════════════════════════════════════════════
h1("켠 뒤에 확인할 것")

shell([
    "systemctl status policy-eval           # 떠 있나",
    "journalctl -u policy-eval -f           # 무슨 일이 있나",
    "curl -s localhost:8000/api/config      # 서버가 답하나",
    "curl -sI https://도메인/records        # 404 여야 한다",
])

p("/api/config 가 이렇게 답하면 정상입니다.")

shell([
    '{"access_required": true,',
    ' "mode": "personal",',
    ' "has_server_key": false,',
    ' "sources": {"fiscal": true, "kdi": true, "opsi": true,',
    '             "bill_key_preset": false}}',
])

table([
    ["보실 것", "이래야 정상"],
    ["mode", "personal — 방문자가 각자 키를 넣는 판"],
    ["access_required", "true — 초대 코드를 물어봅니다"],
    ["has_server_key", "false — 서버에 기관 키가 없습니다.\n"
     "true 면 방문자가 키를 안 넣어도 연구소 키로 돌아갑니다"],
    ["sources 셋", "전부 true — 데이터가 제자리에 있습니다"],
], [2100, 6100])

h2("사람 손으로 한 번")

p("브라우저에서 도메인을 열고 접속 코드를 넣은 뒤, 짧은 아이디어로 "
  "문답을 한 번 끝까지 돌려 보십시오. Claude 키가 필요하므로 시험용으로 "
  "하나 준비하셔야 합니다.")

p("문답이 끝나고 점수와 커버리지가 나오면 네 통로가 모두 살아 있는 "
  "것입니다.")

# ══ 7. 운영 ════════════════════════════════════════════════════════════
h1("운영")

h2("설정을 바꾸려면")

p("전부 /etc/policy-eval.env 한 곳에 있습니다. 고친 뒤 서비스를 다시 "
  "켜면 됩니다.")

shell([
    "sudo vi /etc/policy-eval.env",
    "sudo systemctl restart policy-eval",
])

table([
    ["", "무엇"],
    ["SOCRATIC_ACCESS_CODE", "접속 코드. 방문자에게 알려 줄 값."],
    ["SOCRATIC_MAX_SESSIONS_PER_DAY",
     "하루 세션 상한. 기본 30, 설치 때 100 으로 올려 둡니다.\n"
     "밖에 열어 두는 동안의 안전판이므로 필요 이상으로\n"
     "올리지 않는 편이 낫습니다."],
    ["ASSEMBLY_KEY",
     "국회 의안 인증키 (선택). 넣어 두면 방문자가 안 넣어도\n"
     "③ 통로가 돕니다. 열린국회정보에서 발급합니다 —\n"
     "공공데이터포털(data.go.kr) 키와 다릅니다."],
    ["ANTHROPIC_API_KEY",
     "experiment 판일 때만 채웁니다. personal 판에서 채우면\n"
     "방문자가 키를 안 넣어도 이 키로 돌아가 요금이 연구소에\n"
     "붙습니다. 비워 두십시오."],
], [2900, 5300])

h2("연구 자료")

rich([("webapp/sessions.db 가 연구 자료입니다. ", True),
      ("문답·채점·선례 판정이 전부 여기 쌓입니다. 백업 대상에 넣으시되, "
       "개인정보가 담길 수 있으므로 배포용 zip 에는 넣지 마십시오.",
       False)])

p("사람이 읽을 형태로 보려면:")
shell(["cd /opt/policy-eval && .venv/bin/python webapp/show_session.py"])

h2("갱신")

p("설치할 때와 같은 스크립트를 다시 돌리시면 됩니다. 이미 쌓인 문답"
  "(webapp/sessions.db)과 설정(/etc/policy-eval.env)은 건드리지 "
  "않고 코드만 바꿔 놓습니다.")

shell([
    "unzip 새로받은.zip && cd 정책아이디어평가",
    "sudo bash deploy/setup_server.sh",
    "sudo systemctl restart policy-eval",
])

# ══ 8. 문제 해결 ═══════════════════════════════════════════════════════
h1("잘 안 될 때")

table([
    ["이런 일이 생기면", "이렇게 하십시오"],
    ["서비스가 안 뜬다",
     "journalctl -u policy-eval -n 50\n"
     "대개 가상환경 경로나 권한 문제입니다."],
    ["502 Bad Gateway",
     "서비스가 죽었거나 포트가 다릅니다.\n"
     "systemctl status policy-eval 부터 보십시오."],
    ["429 Too Many Requests",
     "nginx 의 요청 제한(분당 30회)에 걸린 것입니다. 고장이 아닙니다.\n"
     "여러 사람이 같은 공인 IP 로 들어오면 걸릴 수 있습니다 —\n"
     "그때는 nginx.conf 의 rate 를 올리십시오."],
    ["504 Gateway Timeout",
     "Claude 호출이 nginx 기본 타임아웃(60초)에 걸린 것입니다.\n"
     "함께 드린 설정에는 300초로 되어 있습니다 —\n"
     "다른 설정을 쓰고 계신지 보십시오."],
    ["화면은 뜨는데 문답이\n시작되지 않는다",
     "방문자가 API 키를 안 넣었거나 키가 거부됐습니다.\n"
     "화면 「Claude 연결」 칸에 이유가 찍힙니다."],
    ["「닿지 못했습니다」",
     "방화벽이 api.anthropic.com 을 막고 있습니다.\n"
     "망 담당자에게 이 주소를 확인하십시오."],
    ["의안 통로가 의심스럽다",
     ".venv/bin/python webapp/check_bill.py\n"
     "키 없음 / 망 차단 / 키 거부 / 정상을 구분합니다."],
    ["Claude 키가 의심스럽다",
     ".venv/bin/python webapp/check_claude.py\n"
     "붙여넣다 섞여든 공백·보이지 않는 문자를\n"
     "몇 번째 자리인지까지 찍어 줍니다."],
    ["하루 상한에 걸렸다",
     "/etc/policy-eval.env 의\n"
     "SOCRATIC_MAX_SESSIONS_PER_DAY 를 올리고 재시작."],
], [2400, 5800])

h2("401 과 400 은 다릅니다")

p("인증 실패(401)는 키 자체의 문제입니다. 잔액 부족과 한도 초과는 401 이 "
  "아니라 400 으로 나옵니다. 401 이면 돈 문제가 아닙니다.")

# ══ 9. 폴더 ════════════════════════════════════════════════════════════
h1("폴더 안에 무엇이 있나")

table([
    ["", ""],
    ["mode.txt", "personal — 이 한 줄이 판을 정합니다"],
    ["deploy/setup_server.sh", "설치. 이것만 실행하면 됩니다"],
    ["deploy/policy-eval.service", "systemd 유닛"],
    ["deploy/nginx.conf", "앞단 설정"],
    ["deploy/README.md", "더 자세한 설명"],
    ["data/", "평가에 쓰는 데이터 3개 (44MB)"],
    ["webapp/", "웹 서버와 화면. sessions.db 도 여기 생깁니다"],
    ["socratic/", "문답·채점 엔진과 프롬프트"],
    ["kdinov/", "KDI 코퍼스 정밀 판정기"],
    ["읽어보세요.txt", "이 문서의 요약본"],
], [2600, 5600])

p("윈도우용 배치 파일(.bat)은 이 판에 들어 있지 않습니다.")

rich([("압축을 풀면 45MB 를 차지합니다. ", True),
      ("zip 은 8.8MB 이지만 데이터가 압축돼 있어서 그렇습니다. 푸는 데 "
       "30초에서 1분쯤 걸립니다.", False)])

out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "서버설치_안내.docx")
doc.save(out)
print("저장 완료 —", out)
