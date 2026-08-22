#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""평가 세션 하나를 Word 보고서로 만든다. 5_보고서.bat 이 부르는 본체.

화면은 스크롤하며 읽는 것이고 보고서는 남겨서 돌려 보는 것이다. 그래서
화면에 있는 것을 그대로 옮기지 않고, 읽는 순서를 다시 짰다.

  종합 판정 → 5문장 요약 → 축 A 세 기준 → 축 B 실질 독창성
  → 강점·보완 → 부록: 대화 전문

지키는 규율 하나. 조회하지 않은 통로를 "0건"으로 적지 않는다.
미실행 / 조회 실패 / N건 을 끝까지 구분한다. 앞의 둘은 '선례가 없다'는
뜻이 아니라 '모른다'는 뜻이고, 보고서에서 둘을 뭉개면 읽는 사람이
없는 근거를 있는 것으로 오해한다.

사용:  5_보고서.bat  를 더블클릭
       python webapp\\make_report.py            가장 최근 세션
       python webapp\\make_report.py --list     세션 목록
       python webapp\\make_report.py <세션ID>
"""
import json
import os
import sqlite3
import sys
from datetime import datetime
from pathlib import Path

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
DB = HERE / "sessions.db"

CRIT_KO = {"originality": "독창성 — 방어력", "practicality": "실용성",
           "acceptance": "수용태도"}
FIVE_ORDER = ["who", "problem", "how", "different", "evidence"]
FIVE_LABEL = {"who": "누가", "problem": "어떤 문제를", "how": "어떻게 푸는가",
              "different": "무엇이 다른가", "evidence": "무엇으로 뒷받침되는가"}
SRC_KO = {"fiscal": "재정", "prism": "KDI 연구", "bill": "국회 의안",
          "overseas": "해외 사례", "knowledge": "모델 지식"}
CH_KO = {"fiscal": "① 재정(집행)", "prism": "② KDI 연구(검토)",
         "bill": "③ 국회 의안(입법)", "overseas": "④ 해외 사례(시행)"}


def say(*a):
    print(*a, flush=True)


def hold():
    if os.name != "nt" or not sys.stdin.isatty():
        return
    try:
        input("\n  Enter 를 누르면 닫힙니다. ")
    except Exception:
        pass


# ── 자료 읽기 ─────────────────────────────────────────────────────────
def conn():
    c = sqlite3.connect(DB)
    c.row_factory = sqlite3.Row
    return c


def load(sid=None):
    """세션 하나를 통째로 읽는다. sid 를 안 주면 '가장 최근에 끝난' 것을 고른다.

    그냥 최신을 고르면 중간에 그만둔 세션이 잡힌다. 문답을 시작만 하고 닫은
    기록이 훨씬 많이 쌓이기 때문이다. 보고서로 낼 만한 것은 채점까지 끝난
    세션이므로 그쪽을 먼저 찾고, 하나도 없을 때만 최신으로 물러선다.
    """
    c = conn()
    if sid:
        row = c.execute("SELECT * FROM sessions WHERE id = ?", (sid,)).fetchone()
    else:
        row = c.execute(
            "SELECT s.* FROM sessions s JOIN evaluations e ON e.session_id = s.id "
            "ORDER BY s.created_at DESC, s.rowid DESC LIMIT 1").fetchone()
        if not row:
            row = c.execute("SELECT * FROM sessions "
                            "ORDER BY created_at DESC, rowid DESC LIMIT 1").fetchone()
    if not row:
        return None
    s = dict(row)
    ev = c.execute("SELECT result, weighted_total FROM evaluations "
                   "WHERE session_id = ?", (s["id"],)).fetchone()
    s["evaluation"] = json.loads(ev["result"]) if ev else None
    s["weighted_total"] = ev["weighted_total"] if ev else None
    s["axis_b"] = json.loads(s["originality"]) if s.get("originality") else None
    s["turns"] = [dict(r) for r in c.execute(
        "SELECT seq, role, stage, content FROM turns WHERE session_id = ? "
        "ORDER BY seq", (s["id"],))]
    return s


def list_sessions():
    c = conn()
    rows = c.execute(
        "SELECT s.id, s.idea, s.status, s.created_at, s.profile, "
        "e.weighted_total AS score FROM sessions s "
        "LEFT JOIN evaluations e ON e.session_id = s.id "
        "ORDER BY s.created_at DESC, s.id DESC").fetchall()
    say()
    say(f"  {'세션ID':<8} {'일시':<17} {'점수':<6} 아이디어")
    say("  " + "-" * 68)
    for r in rows:
        sc = f"{r['score']:.1f}" if r["score"] is not None else r["status"] or "-"
        idea = (r["idea"] or "").replace("\n", " ")
        say(f"  {str(r['id']):<8} {(r['created_at'] or '')[:16]:<17} {sc:<6} {idea[:38]}")
    say()


# ── 커버리지: 미실행 / 조회 실패 / N건 ────────────────────────────────
def coverage_rows(b):
    """(통로, 상태, 질의어) 목록. 조회하지 않은 것을 0건으로 적지 않는다."""
    lk = (b or {}).get("lookup")
    if not lk:
        return None
    q = lk.get("queries") or {}
    pf = lk.get("profile") or {}
    failed = lk.get("failed") or {}
    # profile 비트가 null 이면 조회기 자체가 돌지 않은 것이다(= 미실행).
    ran = {"fiscal": pf.get("exec") is not None, "prism": pf.get("review") is not None,
           "bill": pf.get("law") is not None, "overseas": pf.get("intl") is not None}
    rows = []
    for src in ("fiscal", "prism", "bill", "overseas"):
        hits = lk.get(src) or []
        if not ran[src]:
            state = "미실행"
        elif failed.get(src):
            state = "조회 실패"
        else:
            state = f"{len(hits)}건"
            if src == "bill" and hits:
                rc = {}
                for x in hits:
                    k = x.get("result") or "계류"
                    rc[k] = rc.get(k, 0) + 1
                state += " (" + ", ".join(f"{k} {v}" for k, v in rc.items()) + ")"
        qs = ", ".join(q.get(src) or []) if ran[src] else "—"
        rows.append((CH_KO[src], state, qs or "—"))
    return rows


# ── Word 만들기 ───────────────────────────────────────────────────────
def build_docx(s, out):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    ev = s["evaluation"]
    b = s["axis_b"]

    def h(text, level=1):
        p = doc.add_heading(text, level=level)
        for r in p.runs:
            r.font.name = "맑은 고딕"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        return p

    def para(text, bold=False, size=10, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = "맑은 고딕"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        return p

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, hd in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            r = cell.paragraphs[0].add_run(hd)
            r.bold = True
            r.font.size = Pt(9)
            r.font.name = "맑은 고딕"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ""
                r = cells[i].paragraphs[0].add_run(str(v))
                r.font.size = Pt(9)
                r.font.name = "맑은 고딕"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        doc.add_paragraph()
        return t

    # ── 표지 ──
    title = doc.add_heading("소크라테스식 아이디어 평가 보고서", level=0)
    for r in title.runs:
        r.font.name = "맑은 고딕"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    p = para(f"세션 {s['id']} · {(s['created_at'] or '')[:16]} · 프로필 {s.get('profile') or '원본'}",
             size=9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    h("평가 대상 아이디어", 1)
    para(s["idea"] or "(기록 없음)")

    # ── 종합 ──
    h("1. 종합 판정", 1)
    if ev:
        para(f"가중 총점  {s['weighted_total']:.1f} / 10", bold=True, size=14)
        if ev.get("score_band"):
            sb = ev["score_band"]
            para(f"평가 범위 {sb['low']} – {sb['high']}  "
                 f"(두 심사위원의 괴리 ±{sb['radius']})", size=9)
        if ev.get("evidence_level"):
            el = ev["evidence_level"]
            para(f"근거 신뢰도 {el['label']} (E{el['level']} · "
                 f"{el['met']}/{el['total']} 충족)", size=9)
        try:
            w = json.loads(s["weights"]) if isinstance(s["weights"], str) else s["weights"]
            para("가중치 — " + " · ".join(
                f"{CRIT_KO.get(k, k)} {v}" for k, v in (w or {}).items()), size=9)
        except Exception:
            pass
    else:
        para("아직 채점되지 않았습니다. 문답만 진행된 상태입니다.", italic=True)

    if b and b.get("band"):
        para(f"실질 독창성  {b['band']}  (확신도 {b.get('confidence', '-')})",
             bold=True, size=12)
        if ev and b["band"] != "판정 보류":
            hi_def = (s["weighted_total"] or 0) >= 6
            hi_org = b["band"] == "계열 밖 시도"
            cell = ("추진" if (hi_def and hi_org) else
                    "기존 사업 개선으로 재포지셔닝" if hi_def else
                    "보완 후 재문답" if hi_org else "보류")
            para(f"위치 — {cell}   (방어 {'높음' if hi_def else '낮음'} · "
                 f"독창성 {'높음' if hi_org else '낮음'})", size=9)
    elif b is None:
        para("실질 독창성(선례 조사)은 수행되지 않았습니다.", italic=True, size=9)

    # ── 5문장 ──
    if ev and ev.get("five_lines"):
        h("2. 다섯 문장 요약", 1)
        rows = [(FIVE_LABEL[k], ev["five_lines"][k])
                for k in FIVE_ORDER if ev["five_lines"].get(k)]
        table(["항목", "내용"], rows)

    # ── 축 A ──
    if ev and ev.get("criteria"):
        h("3. 축 A — 방어력 (12문답 기반)", 1)
        para("각 기준은 체크리스트 10항목과 종합판단을 따로 매긴 뒤 합칩니다. "
             "두 점수의 괴리가 3점 이상이면 평균이 아니라 낮은 쪽을 채택합니다 — "
             "확신이 갈릴 때 후하게 주지 않기 위해서입니다.", size=9, italic=True)
        for c in ev["criteria"]:
            gap = abs(c["checklist_total"] - c["holistic_score"])
            h(f"{CRIT_KO.get(c['key'], c['label'])}  —  {c['final']}/10 "
              f"(가중치 {c['weight']})", 2)
            para(f"체크리스트 {c['checklist_total']}/10 · 종합판단 {c['holistic_score']}/10 → "
                 + (f"괴리 {gap}점 → 낮은 쪽 채택 {c['final']}" if gap >= 3
                    else f"평균 {c['final']}"), size=9)
            para("종합판단 근거: " + c.get("holistic_rationale", ""), size=9)
            table(["", "항목", "판정 근거"],
                  [("충족" if it.get("met") else "미충족", it.get("label", ""),
                    it.get("evidence", "")) for it in c.get("items", [])])

    # ── 축 B ──
    if b:
        h("4. 축 B — 실질 독창성 (선례 조사)", 1)
        para(f"판정 {b.get('band', '-')} · 확신도 {b.get('confidence', '-')}"
             + (f" · 정책 유형 {b['policy_type']}" if b.get("policy_type") else ""),
             bold=True)
        if b.get("reasoning"):
            para("판정 근거: " + b["reasoning"], size=9)

        if b.get("evidence"):
            h("근거", 2)
            table(["등급", "출처", "내용"],
                  [(f"{e.get('grade', '-')}급", SRC_KO.get(e.get("source"), e.get("source") or "-"),
                    e.get("text", "")) for e in b["evidence"]])

        h("검증 통로 커버리지", 2)
        rows = coverage_rows(b)
        if rows:
            table(["통로", "결과", "질의어"], rows)
            para("‘미실행’은 조회기가 돌지 않은 것이고 ‘조회 실패’는 돌았으나 "
                 "오류로 결과를 받지 못한 것입니다. 둘 다 ‘선례가 없다’는 뜻이 "
                 "아니라 ‘확인하지 못했다’는 뜻이며, 판정 근거로 쓰이지 않습니다.",
                 size=9, italic=True)
            lk = b.get("lookup") or {}
            pris = lk.get("prism") or []
            if pris:
                para("KDI 연구: " + "; ".join(
                    x.get("title", "") + (f" [{x.get('code')}·{x.get('role', '')}]"
                                          if x.get("code") else "") for x in pris), size=9)
            over = lk.get("overseas") or []
            if over:
                para("해외 사례(OPSI): " + "; ".join(
                    x.get("title", "") + (f" ({x.get('country')}{' ' + str(x.get('year')) if x.get('year') else ''})"
                                          if x.get("country") else "") for x in over), size=9)
            bills = lk.get("bill") or []
            if bills:
                para("국회 의안: " + "; ".join(
                    x.get("name", "") + (f" [{x.get('result')}]" if x.get("result") else "")
                    for x in bills), size=9)
        else:
            para("외부 조회를 하지 않았습니다 — 모델 지식만으로 판정한 결과입니다. "
                 "선례가 없다는 근거로 쓸 수 없습니다.", italic=True)

        if b.get("claimed_precedents"):
            h("제안자가 지목한 선례 (점수 미반영)", 2)
            table(["지목", "확인"],
                  [(c.get("text", str(c)) if isinstance(c, dict) else str(c),
                    (c.get("verdict", "") if isinstance(c, dict) else ""))
                   for c in b["claimed_precedents"]])

    # ── 강점·보완 ──
    if ev:
        h("5. 강점과 보완", 1)
        for label, key in (("강점", "strengths"), ("보완 제안", "suggestions")):
            if ev.get(key):
                h(label, 2)
                for x in ev[key]:
                    doc.add_paragraph(str(x), style="List Bullet")
        if ev.get("encouragement"):
            h("총평", 2)
            para(ev["encouragement"])

    # ── 부록: 대화 전문 ──
    doc.add_page_break()
    h("부록 — 대화 전문", 1)
    para("평가의 모든 근거는 아래 문답에서 나옵니다. 점수에 의문이 들면 "
         "해당 단계의 문답을 확인하십시오.", size=9, italic=True)
    for t in s["turns"]:
        who = "제안자" if t["role"] in ("user", "proposer") else "질문자"
        head = f"턴 {t['seq']} · {who}" + (f" ({t['stage']})" if t.get("stage") else "")
        para(head, bold=True, size=9)
        para(t["content"] or "")
        doc.add_paragraph()

    doc.save(out)


def main():
    if not DB.exists():
        say(f"\n  기록이 없습니다: {DB}")
        say("  아직 평가를 한 번도 끝내지 않았거나, 다른 폴더의 기록입니다.\n")
        return 1

    args = [a for a in sys.argv[1:]]
    if args and args[0] in ("--list", "-l"):
        list_sessions()
        return 0

    s = load(args[0] if args else None)
    if not s:
        say("\n  그런 세션이 없습니다. 목록:  python webapp\\make_report.py --list\n")
        return 1

    say()
    say("=" * 62)
    say("  소크라테스식 아이디어 평가 — 보고서 만들기")
    say("=" * 62)
    say()
    say(f"  세션   : {s['id']}  ({(s['created_at'] or '')[:16]})")
    say(f"  아이디어: {(s['idea'] or '')[:50]}")
    say(f"  문답   : {len(s['turns'])}턴")
    say(f"  채점   : " + (f"가중 총점 {s['weighted_total']:.1f}"
                          if s["evaluation"] else "아직 채점되지 않음"))
    say(f"  선례조사: " + (s["axis_b"].get("band", "-") if s["axis_b"] else "미실행"))
    say()

    out = ROOT / f"평가보고서_{s['id']}_{datetime.now():%Y%m%d}.docx"
    try:
        build_docx(s, out)
    except ImportError:
        say("  python-docx 가 없어 Word 로 만들지 못했습니다.")
        say("  1_설치.bat 을 다시 실행하거나, 명령창에서:")
        say("      pip install python-docx")
        say()
        return 1
    say(f"  만들었습니다:\n    {out}")
    say()
    say("  더블클릭하면 Word 로 열립니다. PDF 가 필요하면 Word 에서")
    say("  [파일] → [다른 이름으로 저장] → PDF 를 고르세요.")
    say()
    say("=" * 62)
    say()
    return 0


if __name__ == "__main__":
    rc = 1
    try:
        rc = main()
    except Exception as e:
        import traceback
        say("\n  보고서를 만들지 못했습니다 — " + repr(e))
        say(traceback.format_exc())
    finally:
        hold()
    sys.exit(rc)
