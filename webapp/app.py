"""소크라테스 아이디어 평가 — 웹 MVP (FastAPI).

실행 (저장소 루트에서):
    pip install -r webapp/requirements.txt   # fastapi·uvicorn·python-docx·pypdf
    claude /login                      # 최초 1회, Pro 구독 로그인 (API 키 불필요)
    python webapp/app.py               # http://localhost:8000
    # 또는: uvicorn webapp.app:app --port 8000
"""

import base64
import html
import io
import json
import os
import re
import sqlite3
import sys
import threading
import urllib.error
import urllib.parse
import urllib.request
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Optional

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import FileResponse, JSONResponse, RedirectResponse
from pydantic import BaseModel

from socratic import engine

try:
    from webapp import db
except ImportError:  # `python webapp/app.py`로 직접 실행한 경우
    import db

STATIC_DIR = Path(__file__).parent / "static"

app = FastAPI(title="소크라테스 아이디어 평가")
db.init()

# 외부 공개(터널) 시 안전장치 — 환경 변수로 설정
ACCESS_CODE = os.environ.get("SOCRATIC_ACCESS_CODE", "").strip()
MAX_SESSIONS_PER_DAY = int(os.environ.get("SOCRATIC_MAX_SESSIONS_PER_DAY", "30"))
# 1이면 루트(/)를 /policy로 리다이렉트 → 연구자에게 수정판만 노출.
POLICY_ONLY = os.environ.get("SOCRATIC_POLICY_ONLY", "").strip() in ("1", "true", "True")
# 선례 조사 축(축 B) — 공공데이터포털 키 하나로 PRISM·국회 의안을 함께 쓴다.
# 재정(세출예산)은 API가 아니라 로컬 정적 파일(data/fiscal.json)이라 키가 필요 없다.
def _clean_key(name):
    """환경변수 키에서 흔한 실수(따옴표·공백·개행)를 제거한다.
    cmd에서 set KEY=\"abc\" 처럼 넣으면 따옴표가 값에 포함되어 400을 유발한다."""
    return os.environ.get(name, "").strip().strip('"').strip("'").strip()


DATA_GO_KR_KEY = _clean_key("DATA_GO_KR_KEY")      # PRISM
# 국회 의안은 열린국회정보(open.assembly.go.kr) 별도 인증키를 쓴다.
ASSEMBLY_KEY = _clean_key("ASSEMBLY_KEY")


@app.exception_handler(RuntimeError)
def runtime_error_handler(request: Request, exc: RuntimeError):
    """엔진 오류(claude CLI 미설치/미로그인 등)를 사용자가 읽을 수 있는 메시지로 반환."""
    return JSONResponse(status_code=502, content={"detail": str(exc)})


class CreateRequest(BaseModel):
    idea: str
    weights: dict | None = None  # {"originality": w1, "practicality": w2, "acceptance": w3}
    access_code: str | None = None
    profile: str | None = None   # 원본 | 정책 (수정판). 없으면 원본.


class AnswerRequest(BaseModel):
    answer: str


def _format_turn(t):
    if t["role"] == "questioner":
        return f"[턴 {t['seq']}] 질문자({t['stage']}): {t['content']}"
    return f"[턴 {t['seq']}] 제안자: {t['content']}"


def _transcript_log(sid):
    return [_format_turn(t) for t in db.get_turns(sid)]


def _next_seq(sid):
    turns = db.get_turns(sid)
    return (turns[-1]["seq"] + 1) if turns else 1


def _progress(stage_index, q_in_stage):
    _, label, n_turns, _ = engine.STAGES[stage_index]
    return {
        "stage_index": stage_index,
        "stage_label": label,
        "question_no": q_in_stage,
        "questions_in_stage": n_turns,
        "total_stages": len(engine.STAGES),
    }


def _ask_and_store(sid, stage_index):
    """현재 단계의 다음 질문을 생성해 저장하고 반환한다."""
    name, label, _, directive = engine.STAGES[stage_index]
    # 정책 프로필의 독창성 라운드에만 선례 지목 유도문을 덧붙인다(원본 무영향).
    session = db.get_session(sid)
    if session and session["profile"] == "정책" and name == "originality":
        directive = directive + " " + engine.PRECEDENT_ANCHOR_LINE
    question = engine.ask_questioner(_transcript_log(sid), directive)
    db.add_turn(sid, _next_seq(sid), "questioner", label, question)
    return question


def _grade_session(sid, weights, profile="원본"):
    """대화 로그를 채점하고 결과를 저장한다."""
    result = engine.grade("\n\n".join(_transcript_log(sid)), profile)
    total = round(engine.weighted_total(result, weights), 2)
    db.save_evaluation(sid, result, total)
    db.set_status(sid, "graded")
    return _evaluation_payload(result, total, weights, profile)


def _evaluation_payload(result, total, weights, profile="원본"):
    labels = engine.checklist_labels(profile)
    payload = {
        "profile": profile,
        "criteria": [
            {
                "key": c,
                "label": engine.CRITERIA_KO[c],
                "weight": weights[c],
                "final": result["criteria"][c]["final"],
                "checklist_total": result["criteria"][c]["checklist"]["total"],
                "holistic_score": result["criteria"][c]["holistic"]["score"],
                "holistic_rationale": result["criteria"][c]["holistic"]["rationale"],
                "items": [
                    {
                        "id": item_id,
                        "label": labels[c][item_id],
                        "met": entry["met"],
                        "evidence": entry["evidence"],
                    }
                    for item_id, entry in result["criteria"][c]["checklist"]["items"].items()
                ],
            }
            for c in engine.CRITERIA
        ],
        "weighted_total": total,
        "strengths": result["strengths"],
        "suggestions": result["suggestions"],
        "encouragement": result["encouragement"],
    }
    # 5문장 프레임·점수 범위·근거신뢰도는 원본·수정판 모두에 얹는다.
    # 전부 코드 계산 — 추가 Claude 호출 없음. (프로필 차이는 A9 채점뿐.)
    payload["five_lines"] = engine.five_lines(result, weights)
    payload["score_band"] = engine.score_band(result, weights)
    payload["evidence_level"] = engine.evidence_level(result)
    return payload


@app.get("/")
def index():
    # 터널 배포(SOCRATIC_POLICY_ONLY=1) 시 원본 대신 수정판으로 보낸다.
    if POLICY_ONLY:
        return RedirectResponse(url="/policy")
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/policy")
def policy():
    """수정판(정책 평가) — 5문장 프레임·점수 범위·강화 A9 채점."""
    return FileResponse(STATIC_DIR / "policy.html")


@app.get("/api/config")
def get_config():
    return {"access_required": bool(ACCESS_CODE)}


@app.post("/api/sessions")
def create_session(req: CreateRequest):
    if ACCESS_CODE and (req.access_code or "").strip() != ACCESS_CODE:
        raise HTTPException(403, "초대 코드가 올바르지 않습니다.")
    if db.count_sessions_today() >= MAX_SESSIONS_PER_DAY:
        raise HTTPException(429, "오늘 사용 가능한 세션 수를 모두 사용했습니다. 내일 다시 시도해 주세요.")
    idea = req.idea.strip()
    if not idea:
        raise HTTPException(400, "아이디어가 비어 있습니다.")
    weights = req.weights or dict(engine.DEFAULT_WEIGHTS)
    if set(weights) != set(engine.CRITERIA):
        raise HTTPException(400, "가중치 키가 올바르지 않습니다.")
    if abs(sum(weights.values()) - 1.0) > 1e-6:
        raise HTTPException(400, f"가중치의 합은 1이어야 합니다 (현재 {sum(weights.values()):.2f})")
    profile = (req.profile or "원본").strip()
    if profile not in engine.PROFILES:
        raise HTTPException(400, "알 수 없는 평가 프로필입니다.")

    sid = db.create_session(idea, weights, profile)
    db.add_turn(sid, 1, "proposer", None, idea)
    question = _ask_and_store(sid, 0)
    db.update_progress(sid, 0, 1)
    return {"session_id": sid, "question": question, "progress": _progress(0, 1)}


@app.post("/api/sessions/{sid}/answer")
def answer(sid: str, req: AnswerRequest):
    session = db.get_session(sid)
    if not session:
        raise HTTPException(404, "세션이 없습니다.")
    if session["status"] != "active":
        raise HTTPException(400, "이미 채점이 끝난 세션입니다.")
    text = req.answer.strip()
    if not text:
        raise HTTPException(400, "답변이 비어 있습니다.")

    db.add_turn(sid, _next_seq(sid), "proposer", None, text)

    stage_index = session["stage_index"]
    q_in_stage = session["q_in_stage"]
    # 현재 단계의 질문을 모두 소화했으면 다음 단계로
    if q_in_stage >= engine.STAGES[stage_index][2]:
        stage_index += 1
        q_in_stage = 0

    if stage_index >= len(engine.STAGES):
        weights = json.loads(session["weights"])
        evaluation = _grade_session(sid, weights, session["profile"])
        return {"done": True, "evaluation": evaluation}

    question = _ask_and_store(sid, stage_index)
    q_in_stage += 1
    db.update_progress(sid, stage_index, q_in_stage)
    return {"done": False, "question": question,
            "progress": _progress(stage_index, q_in_stage)}


@app.post("/api/sessions/{sid}/finish")
def finish(sid: str):
    """남은 문답을 건너뛰고 지금까지의 로그로 바로 채점한다."""
    session = db.get_session(sid)
    if not session:
        raise HTTPException(404, "세션이 없습니다.")
    if session["status"] != "active":
        raise HTTPException(400, "이미 채점이 끝난 세션입니다.")
    weights = json.loads(session["weights"])
    evaluation = _grade_session(sid, weights, session["profile"])
    return {"done": True, "evaluation": evaluation}


@app.get("/api/sessions/{sid}")
def get_session(sid: str):
    session = db.get_session(sid)
    if not session:
        raise HTTPException(404, "세션이 없습니다.")
    weights = json.loads(session["weights"])
    payload = {
        "session_id": sid,
        "idea": session["idea"],
        "status": session["status"],
        "weights": weights,
        "turns": db.get_turns(sid),
    }
    ev = db.get_evaluation(sid)
    if ev:
        payload["evaluation"] = _evaluation_payload(
            ev["result"], ev["weighted_total"], weights, session["profile"]
        )
    return payload


# ── 선례 조사 축(축 B) ────────────────────────────────────────────────────
# 세 소스는 서로 다른 질문에 답한다: 재정(집행)·PRISM(검토)·국회 의안(입법).
# 재정은 로컬 정적 파일(연 1회 갱신)이라 키가 없다. PRISM은 DATA_GO_KR_KEY로
# apis.data.go.kr를, 국회 의안은 ASSEMBLY_KEY로 열린국회정보(open.assembly.go.kr,
# ALLBILLV2)를 호출한다. 실패/키 없음이면 해당 소스를 건너뛰고, 미조회는
# 미발견으로 처리하지 않는다(profile 비트에서 None → 화면 '-').
PRISM_BASE = os.environ.get(
    "PRISM_BASE", "https://apis.data.go.kr/1741000/prism_v2/getResearchList_v2")
# 국회 의안: 열린국회정보 '의안정보 통합 API'(ALLBILLV2). ASSEMBLY_KEY를 쓴다.
# 필수 파라미터: KEY·Type·pIndex·pSize + ERACO(대수). 검색은 BILL_NM(의안명).
ALLBILL_BASE = os.environ.get(
    "ALLBILL_BASE", "https://open.assembly.go.kr/portal/openapi/ALLBILLV2")
# ERACO(대수)가 필수라 대수별로 조회한다. 최근 대수 위주(연구 목적상 최근 선례가 중요).
ERACO_TERMS = [t.strip() for t in os.environ.get(
    "ERACO_TERMS", "제22대,제21대").split(",") if t.strip()]
# 페이지 크기. 일부 열린국회 키는 5로 제한되어 100을 보내면 400이 난다(브라우저 확인).
# 상위 몇 건만 쓰므로 5로 충분. 전체 승인 키를 쓰면 BILL_PSIZE로 올릴 수 있다.
BILL_PSIZE = int(os.environ.get("BILL_PSIZE", "5"))
# 제안이유·주요내용 본문: ALLBILLV2엔 없다. 의안정보시스템(likms)의 요약 팝업에서
# BILL_ID로 받아온다(인증키 불필요, 공개 웹). 상위 5건만 조회한다.
# 새 의안정보시스템: 상세 페이지(껍데기)는 billDetailPage.do, 제안이유 본문은 그
# 안에서 billInfo.do가 불러온다(DevTools로 확인). billInfo.do를 직접 호출한다.
LIKMS_DETAIL_BASE = os.environ.get(
    "LIKMS_DETAIL_BASE", "https://likms.assembly.go.kr/bill/bi/billDetailPage.do")
LIKMS_BILLINFO = os.environ.get(
    "LIKMS_BILLINFO", "https://likms.assembly.go.kr/bill/bi/bill/detail/billInfo.do")
LIKMS_MENU_NO = os.environ.get("LIKMS_MENU_NO", "2600044")
BILL_SUMMARY_MAXLEN = int(os.environ.get("BILL_SUMMARY_MAXLEN", "500"))  # 본문 앞에서 500자
FISCAL_JSON = ROOT / "data" / "fiscal.json"
_TIMEOUT = 8
_fiscal_cache = None


def _redact(url):
    """로그용: 인증키 값을 가린다(경로·나머지 파라미터는 보이게)."""
    return re.sub(r"([?&](?:serviceKey|ServiceKey|KEY)=)[^&]+", r"\1***", url)


# 일부 정부 API 방화벽이 Python-urllib 기본 User-Agent를 400으로 막는다.
# 브라우저처럼 보이게 헤더를 붙인다(브라우저로는 되는데 서버로만 400인 원인).
_UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
       "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36")


def _urlopen_read(url, accept, referer=None, data=None, timeout=None):
    """공통 GET/POST. HTTPError면 응답 본문(진짜 사유)까지 담아 RuntimeError로 올린다.
    referer를 주면 헤더에 넣는다(likms처럼 Referer를 검사하는 사이트용).
    data(bytes)를 주면 POST로 보낸다(폼 인코딩). timeout으로 대기시간을 늘릴 수 있다."""
    headers = {"Accept": accept, "User-Agent": _UA, "Accept-Language": "ko,en;q=0.9"}
    if referer:
        headers["Referer"] = referer
    if data is not None:
        headers["Content-Type"] = "application/x-www-form-urlencoded; charset=UTF-8"
        headers["X-Requested-With"] = "XMLHttpRequest"
    req = urllib.request.Request(url, data=data, headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=timeout or _TIMEOUT) as resp:
            return resp.read().decode("utf-8", "replace")
    except urllib.error.HTTPError as e:
        body = ""
        try:
            body = e.read().decode("utf-8", "replace")
        except Exception:
            pass
        raise RuntimeError(
            f"HTTP {e.code} {e.reason} | URL {_redact(url)} | 본문 {body[:400]}")


def _http_get_json(url, timeout=None):
    return json.loads(_urlopen_read(url, "application/json", timeout=timeout))


def _http_get_data(url):
    """JSON이면 json, XML이면 dict/list로 변환해 반환한다. data.go.kr 일부
    오퍼레이션은 XML만 주므로, 어느 쪽이 와도 _find_key·_pick으로 다루게 만든다."""
    raw = _urlopen_read(url, "application/json, application/xml")
    s = raw.lstrip("﻿ \t\r\n")
    if s[:1] in ("{", "["):
        return json.loads(raw)
    try:
        return _xml_to_obj(ET.fromstring(raw))
    except ET.ParseError:
        return {}


def _xml_to_obj(el):
    """XML 엘리먼트를 중첩 dict/list로. 같은 태그 반복은 리스트로 묶는다."""
    kids = list(el)
    if not kids:
        return (el.text or "").strip()
    out = {}
    for c in kids:
        tag = c.tag.split("}")[-1]  # 네임스페이스 접두사 제거
        v = _xml_to_obj(c)
        if tag in out:
            if not isinstance(out[tag], list):
                out[tag] = [out[tag]]
            out[tag].append(v)
        else:
            out[tag] = v
    return out


def _find_rows(obj):
    """중첩 JSON에서 dict들의 리스트(레코드 집합)를 재귀로 찾는다."""
    if isinstance(obj, list):
        if obj and all(isinstance(x, dict) for x in obj):
            return obj
        for x in obj:
            r = _find_rows(x)
            if r:
                return r
    elif isinstance(obj, dict):
        for v in obj.values():
            r = _find_rows(v)
            if r:
                return r
    return []


def _pick(row, *keys):
    for k in keys:
        if k in row and row[k] not in (None, ""):
            return row[k]
    return None


def _find_key(obj, key):
    """중첩 JSON에서 특정 키의 값을 재귀로 찾는다(첫 번째)."""
    if isinstance(obj, dict):
        if key in obj:
            return obj[key]
        for v in obj.values():
            r = _find_key(v, key)
            if r is not None:
                return r
    elif isinstance(obj, list):
        for x in obj:
            r = _find_key(x, key)
            if r is not None:
                return r
    return None


def _as_rows(node):
    """레코드 컨테이너를 dict 리스트로 정규화(단건이면 리스트로 감싼다)."""
    if isinstance(node, list):
        return [x for x in node if isinstance(x, dict)]
    if isinstance(node, dict):
        return [node]
    return []


# ── 재정: 로컬 정적 파일 검색 (API 아님) ──
def _fiscal_available():
    return FISCAL_JSON.exists()


def _load_fiscal():
    global _fiscal_cache
    if _fiscal_cache is None:
        try:
            data = json.loads(FISCAL_JSON.read_text(encoding="utf-8"))
            # {unit, items:[...]} 래퍼 또는 [...] 배열 모두 허용.
            _fiscal_cache = data["items"] if isinstance(data, dict) else data
        except Exception:
            _fiscal_cache = []
    return _fiscal_cache


# 세 소스 공용 키워드 매칭. 흔한 일반어 단독 매칭을 막아 거짓 양성을 줄인다.
# 규칙: 전체 문구가 통째로 들어있거나, '의미 있는' 토큰이 2개 이상 함께 겹칠 때만 히트.
# (의미 있는 토큰이 하나뿐인 질의는 그 하나만 겹쳐도 히트 — 과소매칭 방지.)
_STOPWORDS = frozenset({
    "분석", "연구", "지원", "방안", "정책", "사업", "효과", "제도", "개선", "강화",
    "관리", "활성화", "촉진", "계획", "전략", "체계", "현황", "실태", "평가", "도입",
    "운영", "구축", "확대", "및", "등", "관한", "위한", "대한", "기반", "관련",
})


def _keyword_hit(query, *texts):
    hay = " ".join(t for t in texts if t)
    q = (query or "").strip()
    if not q or not hay:
        return False
    if q in hay:
        return True
    meaningful = [t for t in q.split() if len(t) >= 2 and t not in _STOPWORDS]
    if not meaningful:                       # 질의가 전부 일반어/단문자면 원래 토큰 사용
        meaningful = [t for t in q.split() if t]
    present = sum(1 for t in meaningful if t in hay)
    return present >= 2 if len(meaningful) >= 2 else present >= 1


def _fiscal_local_search(query):
    """세부사업명 매칭(공용 규칙). 히트 과다 시 예산액 큰 순 5건."""
    q = (query or "").strip()
    if not q:
        return []
    out = [rec for rec in _load_fiscal() if _keyword_hit(q, rec.get("name", ""))]
    out.sort(key=lambda r: max((s.get("amount", 0) for s in r.get("series", [])),
                               default=0), reverse=True)
    return out[:5]


# ── PRISM: 정책연구 과제 (API) ──
# getResearchList_v2의 요청 파라미터는 serviceKey·organ_id·start_date·end_date·
# numOfRows·pageNo 뿐이다(요청변수 표 확인). 즉 '주제 키워드 검색'을 지원하지 않고
# 기관코드·날짜로만 거른다. 그래서 날짜 범위로 여러 페이지를 받아 과제명·사업명에
# 주제어가 있는 것만 로컬에서 골라낸다(최선의 근사; 최근 창 밖 연구는 못 잡을 수 있음).
PRISM_START = os.environ.get("PRISM_START", "20180101")
PRISM_END = os.environ.get("PRISM_END", "20261231")
PRISM_TIMEOUT = int(os.environ.get("PRISM_TIMEOUT", "25"))  # PRISM API가 느려 별도 대기
PRISM_ROWS = int(os.environ.get("PRISM_ROWS", "100"))       # 페이지당 결과 수
PRISM_PAGES = int(os.environ.get("PRISM_PAGES", "3"))       # 훑을 페이지 수(로컬 필터)


def _decode_key(key):
    """data.go.kr 키를 raw로 정규화한다. Encoding 키(%2B·%2F·%3D 포함)든 Decoding
    키든 넣을 수 있게 unquote로 통일 → urlencode가 다시 정확히 인코딩한다."""
    return urllib.parse.unquote(key or "")


def _prism_search_terms(query):
    """로컬 필터에 쓰는 '구체적 주제어'(일반어 제외). PRISM은 키워드 검색이 없어
    이 말들이 과제명·사업명에 있는지로만 관련성을 판정한다(진단 스크립트도 이걸 씀)."""
    toks = _bill_distinct_tokens(query)
    if not toks:
        toks = [t for t in (query or "").split() if len(t) >= 2 and t not in _STOPWORDS]
    return sorted(dict.fromkeys(toks), key=len, reverse=True)[:2]


def _prism_lookup(query):
    if not DATA_GO_KR_KEY:
        return []
    q = (query or "").strip()
    distinct = _bill_distinct_tokens(q) or [
        t for t in q.split() if len(t) >= 2 and t not in _STOPWORDS]
    if not distinct:
        return []
    out, seen = [], set()
    for page in range(1, PRISM_PAGES + 1):
        try:
            params = {"serviceKey": _decode_key(DATA_GO_KR_KEY), "type": "json",
                      "start_date": PRISM_START, "end_date": PRISM_END,
                      "numOfRows": PRISM_ROWS, "pageNo": page}
            data = _http_get_json(PRISM_BASE + "?" + urllib.parse.urlencode(params),
                                  timeout=PRISM_TIMEOUT)
            rows = _as_rows(_find_key(data, "research"))
            if not rows:
                if page == 1:
                    _debug_once("prism", data)
                break
            for r in rows:
                # 확정 필드: research_name(과제명)·organ_name(기관)·research_date(기간).
                title = _pick(r, "research_name", "biz_name")
                if not title or title in seen:
                    continue
                hay = f"{_pick(r, 'research_name') or ''} {_pick(r, 'biz_name') or ''}"
                if not any(t in hay for t in distinct):
                    continue             # 주제어가 과제명에 없으면 무관 → 제외
                seen.add(title)
                out.append({"title": title,
                            "org": _pick(r, "organ_name"),
                            "period": _pick(r, "research_date")})
            if len(out) >= 5 or len(rows) < PRISM_ROWS:
                break                    # 5건 채웠거나 마지막 페이지면 중단
        except Exception as e:
            print(f"prism lookup 실패(p{page}):", e, file=sys.stderr)
            break
    return out[:5]


# ── 국회 의안: 열린국회정보 '의안정보 통합 API'(ALLBILLV2), ASSEMBLY_KEY ──
# 필수: KEY·Type·pIndex·pSize + ERACO(대수). 검색은 BILL_NM(의안명, 서버측 부분일치).
# 응답은 열린국회 표준({"ALLBILLV2":[{head},{row}]})이라 row 컨테이너를 꺼낸다.
# ERACO가 필수라 최근 대수별로 조회해 합친다. 목록 API라 제안이유 본문은 없을 수
# 있다 — 본문 필드가 있으면 summary에 담고(있는 데까지), 없으면 제목·결과만.
_DEBUG_SEEN = set()


def _debug_once(tag, obj):
    """미지수 확정용: 처음 한 번만 응답 구조를 콘솔에 찍는다(복사해 주면 필드 확정)."""
    if tag in _DEBUG_SEEN:
        return
    _DEBUG_SEEN.add(tag)
    try:
        preview = json.dumps(obj, ensure_ascii=False)[:800]
    except Exception:
        preview = str(obj)[:800]
    print(f"[축B 디버그:{tag}] {preview}", file=sys.stderr)


def _strip_html(s):
    """HTML을 순수 텍스트로. script/style 제거, 태그 제거, 엔티티 복원, 공백 정리."""
    s = re.sub(r"(?is)<(script|style)\b.*?</\1>", " ", s)
    s = re.sub(r"(?s)<[^>]+>", " ", s)
    return re.sub(r"\s+", " ", html.unescape(s)).strip()


def _bill_summary(bill_id):
    """2단계: BILL_ID로 의안정보시스템(likms)에서 제안이유·주요내용 본문을 받는다.

    ※ 미완(나중에 보완): 신형 likms는 SPA라 billInfo.do가 섹션별로 조각을 주는데,
      제안이유 섹션을 선택하는 정확한 파라미터를 아직 못 잡아(POST는 '심사진행'만 옴)
      대개 빈 값이 된다. 실패해도 무해하며(제목·결과만 남음), 향후 그 파라미터를
      billInfo.do의 실제 요청(Copy as cURL)으로 확정하면 본문까지 살릴 수 있다.
      코드는 그 확정을 쉽게 하도록 GET→POST·Referer·디버그 로깅을 남겨둔다."""
    if not bill_id:
        return ""
    referer = LIKMS_DETAIL_BASE + "?" + urllib.parse.urlencode(
        {"billId": bill_id, "currMenuNo": LIKMS_MENU_NO})
    qs = urllib.parse.urlencode({"billId": bill_id, "currMenuNo": LIKMS_MENU_NO})
    # billInfo.do를 GET(?billId=) → POST(billId=) 순으로 시도한다(요청 방식 불명).
    # '제안이유'가 실제로 있는 응답만 본문으로 인정한다(404·껍데기 페이지 배제).
    for method, url, data in (("GET", LIKMS_BILLINFO + "?" + qs, None),
                              ("POST", LIKMS_BILLINFO, qs.encode("utf-8"))):
        try:
            text = _strip_html(_urlopen_read(url, "text/html", referer=referer, data=data))
            _debug_once(f"bill-summary-{method}", text[:500])
            m = re.search(r"제안이유", text)
            if m:
                body = text[m.start():].strip()
                if len(body) >= 40:
                    return body[:BILL_SUMMARY_MAXLEN]
        except Exception as e:
            print(f"bill summary(likms {method}) 실패:", e, file=sys.stderr)
    return ""


# 의안 검색·매칭에서 빼는 '너무 흔한' 도메인 일반어(정책·교육 행정 어휘).
# 이런 말로 의안명을 검색하면 무관한 개정안이 쏟아지므로 구체적 주제어만 쓴다.
_BILL_BROAD = _STOPWORDS | frozenset({
    "교육", "학교", "학생", "수업", "교과", "과정", "교육과정", "신설", "지정",
    "필수", "의무", "의무화", "시행", "국가", "국민", "서비스", "정보", "프로그램",
    "지자체", "활동", "확보", "마련", "추진", "조성",
})


def _is_lawname(tok):
    return tok.endswith(("법", "법률", "법안", "법률안"))


def _bill_distinct_tokens(query):
    """의안 관련성 판정용 '구체적 주제어'. 법률명·일반어를 뺀 고유 명사류."""
    toks = [t for t in (query or "").split()
            if len(t) >= 2 and t not in _BILL_BROAD and not _is_lawname(t)]
    return list(dict.fromkeys(toks))


def _bill_search_terms(query):
    """열린국회 BILL_NM 검색용 핵심어. 법률명(개정안 폭주)·일반어를 빼고 구체적
    주제어를 우선 쓴다. 없으면 일반 의미어라도 사용(최대 2개)."""
    toks = _bill_distinct_tokens(query)
    if not toks:
        toks = [t for t in (query or "").split() if len(t) >= 2 and t not in _STOPWORDS]
    return sorted(dict.fromkeys(toks), key=len, reverse=True)[:2]


def _bill_lookup(query):
    """짧은 핵심어로 ALLBILLV2를 대수별 조회해 후보를 모으고, 각 후보에 likms
    제안이유 본문을 붙인 뒤, 원 질의어와 '이름+본문'으로 정밀 필터링한다.
    (의안명은 법률명이라 제목만으로는 주제를 못 담으므로 본문까지 보고 판정한다.)"""
    if not ASSEMBLY_KEY:
        return []
    q = (query or "").strip()
    terms = _bill_search_terms(q)
    if not terms:
        return []
    distinct = _bill_distinct_tokens(q)
    cand, seen = [], set()
    for term in terms:
        for eraco in ERACO_TERMS:
            try:
                params = {"KEY": ASSEMBLY_KEY, "Type": "json", "pIndex": 1,
                          "pSize": BILL_PSIZE, "ERACO": eraco, "BILL_NM": term}
                data = _http_get_data(ALLBILL_BASE + "?" + urllib.parse.urlencode(params))
                rows = _as_rows(_find_key(data, "row"))
                if not rows:
                    _debug_once(f"bill-{term}-{eraco}", data)
                    continue
                for r in rows:
                    name = _pick(r, "BILL_NM", "BILL_NAME")
                    if name and name not in seen:
                        seen.add(name)
                        cand.append(r)
            except Exception as e:
                print(f"allbillv2 lookup 실패({term}/{eraco}):", e, file=sys.stderr)
    out = []
    for r in cand[:12]:                     # 본문 조회 상한(호출 최소화)
        name = _pick(r, "BILL_NM", "BILL_NAME")
        body = _bill_summary(_pick(r, "BILL_ID", "BILL_NO"))
        hay = f"{name} {body}"
        # 구체적 주제어가 하나라도 이름·본문에 있으면 관련 의안으로 본다(등급기가
        # 설계 차이를 판단). 구체어가 없는 질의면 공용 규칙으로 보수적으로 판정.
        if distinct:
            if not any(t in hay for t in distinct):
                continue
        elif not _keyword_hit(q, name, body):
            continue
        out.append({
            "name": name,
            "proposer": _pick(r, "PROPOSER", "PPSR_NM", "RPPSR_NM"),
            "date": _pick(r, "PPSL_DT", "PROPOSE_DT", "PPSL_DATE"),
            "committee": _pick(r, "JRCMIT_NM", "CURR_COMMITTEE", "COMMITTEE_NM"),
            # 본회의 결과 → 소관위 결과 → 처리단계 → 계류구분 순으로 가장 구체적인 값.
            "result": _pick(r, "RGS_CONF_RSLT", "JRCMIT_PROC_RSLT", "PROC_STAGE_CD",
                            "PASSGUBN") or "계류",
            "summary": body,
            "link": _pick(r, "LINK_URL", "DETAIL_LINK"),
            "eraco": _pick(r, "ERACO"),
        })
        if len(out) >= 5:
            break
    return out


class LookupRequest(BaseModel):
    query: str


@app.post("/api/prism")
def api_prism(req: LookupRequest):
    return {"hits": _prism_lookup(req.query.strip())}


@app.post("/api/bill")
def api_bill(req: LookupRequest):
    return {"hits": _bill_lookup(req.query.strip())}


# ── 검토(연구) 축: KDI 정책연구 로컬 코퍼스(kdi/kdi_corpus.db) ──
# PRISM(API)이 불안정해 기본 정지. reports 테이블이 채워지면 KDI가 자동 활성화된다.
# PRISM을 되살리려면 환경변수 PRISM_ENABLED=1.
PRISM_ENABLED = os.environ.get("PRISM_ENABLED", "0").strip().lower() in ("1", "true", "yes")
KDI_DB = ROOT / "kdi" / "kdi_corpus.db"

_KDI_STOP = {"및", "등", "관한", "관련", "대한", "위한", "통한", "그리고", "또는",
             "the", "and", "for", "with", "of", "in", "on", "to", "study", "연구",
             "정책", "방안", "분석", "제도", "개선", "방향", "과제"}


# kdinov(정교한 KDI 독창성 판정기) 연동 — docs 스키마 코퍼스(kdi.sqlite)를 쓴다.
# 있으면 kdinov(분해→검색→2차원 판정)를, 없으면 위 naive reports 조회를 쓴다.
KDI_SQLITE = Path(os.environ.get("KDI_SQLITE", str(ROOT / "kdi" / "kdi.sqlite")))
_KDINOV = None
_KDI_CORPUS = None


def _kdinov():
    """kdinov 모듈 묶음(없으면 None). 한 번만 로드."""
    global _KDINOV
    if _KDINOV is None:
        try:
            from kdinov import sources as _s, model as _m, decompose as _d
            from kdinov import search as _se, verdict as _v
            _KDINOV = {"sources": _s, "model": _m, "decompose": _d,
                       "search": _se, "verdict": _v}
        except Exception as e:
            print("kdinov 로드 실패:", e, file=sys.stderr)
            _KDINOV = False
    return _KDINOV or None


def _kdi_corpus():
    """kdinov docs 코퍼스(list[Doc]) 캐시. 없으면 빈 리스트."""
    global _KDI_CORPUS
    if _KDI_CORPUS is None:
        k = _kdinov()
        if k and KDI_SQLITE.exists():
            try:
                _KDI_CORPUS = k["sources"].Store(str(KDI_SQLITE)).all()
                print(f"kdinov 코퍼스 로드: {len(_KDI_CORPUS)}건 · {KDI_SQLITE}", file=sys.stderr)
            except Exception as e:
                print("kdi corpus 로드 실패:", e, file=sys.stderr)
                _KDI_CORPUS = []
        else:
            _KDI_CORPUS = []
    return _KDI_CORPUS


def _kdinov_lookup(query):
    """kdinov로 KDI 코퍼스 대조: decompose→search→assess. 상위 5건에 code/role 부착.
    kdinov/코퍼스가 없으면 None(→naive 폴백)."""
    k = _kdinov()
    corpus = _kdi_corpus()
    if not k or not corpus:
        return None
    try:
        idea = k["model"].Idea.from_dict(k["decompose"].decompose_policy_idea(query or ""))
        hits = k["search"].search_docs(corpus, k["search"].terms_from_idea(idea), limit=5)
        out = []
        for h in hits:
            d = h.doc
            a = k["verdict"].assess(d, idea)
            out.append({
                "title": d.title, "org": d.kind or d.source, "period": d.year(),
                "summary": (h.snippet or d.summary or "")[:500], "url": d.url,
                "code": a.code, "role": a.role, "score": a.score,
            })
        return out
    except Exception as e:
        print("kdinov lookup 실패:", e, file=sys.stderr)
        return None


def _kdi_available():
    if _kdi_corpus():                      # kdinov docs 코퍼스가 있으면 활성
        return True
    if not KDI_DB.exists():                # 아니면 naive reports 코퍼스
        return False
    try:
        with sqlite3.connect(KDI_DB) as c:
            return c.execute("SELECT COUNT(*) FROM reports").fetchone()[0] > 0
    except Exception:
        return False


def _kdi_lookup(query):
    """검토(연구) 슬롯: kdinov 우선, 없으면 naive reports 조회."""
    hits = _kdinov_lookup(query)
    if hits is not None:
        return hits
    return _kdi_naive_lookup(query)


def _kdi_naive_lookup(query):
    """폴백: KDI reports 코퍼스에서 제목·본문·키워드를 LIKE 검색해 상위 5건.
    반환 형식은 PRISM과 동일({title, org, period}) + summary."""
    raw = [w.lower() for w in re.findall(r"[0-9A-Za-z가-힣][\w가-힣\-]{1,}", query or "")]
    toks = [t for t in raw if t not in _KDI_STOP and len(t) >= 2] or raw
    if not toks or not KDI_DB.exists():
        return []
    try:
        conn = sqlite3.connect(KDI_DB)
        conn.row_factory = sqlite3.Row
        where = " OR ".join(["cleaned_content LIKE ? OR title LIKE ? OR keywords LIKE ?"] * len(toks))
        params = []
        for t in toks:
            params += [f"%{t}%", f"%{t}%", f"%{t}%"]
        rows = [dict(r) for r in conn.execute(
            f"SELECT * FROM reports WHERE {where} LIMIT 120", params).fetchall()]
        conn.close()
    except Exception as e:
        print("kdi lookup 실패:", e, file=sys.stderr)
        return []
    scored = []
    for d in rows:
        title = (d.get("title") or "").lower()
        hay = f"{title} {(d.get('keywords') or '').lower()} {(d.get('cleaned_content') or '').lower()}"
        score = 0
        for t in toks:
            if t in title:
                score += 3
            elif t in hay:
                score += 1
        for a, b in zip(toks, toks[1:]):
            if f"{a} {b}" in hay:
                score += 3
        if score:
            scored.append((score, d))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [{
        "title": d.get("title"), "org": d.get("org") or d.get("organization"),
        "period": d.get("period") or d.get("year"),
        "summary": (d.get("cleaned_content") or "")[:500], "url": d.get("url") or d.get("source_url"),
    } for _, d in scored[:5]]


# ── 해외 축: OPSI 로컬 DB(overseas/opsi_policies.db) 검색 ──
# 사람이 Claude in Chrome으로 수집·임포트하면 채워진다. 비어 있으면 자동으로 꺼짐.
OPSI_DB = ROOT / "overseas" / "opsi_policies.db"


def _opsi_available():
    if not OPSI_DB.exists():
        return False
    try:
        with sqlite3.connect(OPSI_DB) as c:
            return c.execute("SELECT COUNT(*) FROM cases").fetchone()[0] > 0
    except Exception:
        return False


# OPSI 사례 거의 전부에 등장해 변별력이 없는 초일반어(검색 잡음 제거).
_OPSI_STOP = {
    "the", "and", "for", "with", "from", "that", "this", "are", "was", "has", "have",
    "not", "but", "all", "any", "can", "will", "new", "use", "using", "used", "based",
    "into", "per", "via", "its", "their", "our", "your", "more", "most", "such", "than",
    "then", "they", "them", "also", "may", "one", "two", "who", "how", "what", "when",
    "where", "which", "while", "been", "being", "were", "would", "could", "should",
    "about", "over", "under", "between", "within", "across", "through",
    "public", "government", "governmental", "service", "services", "sector", "innovation",
    "innovative", "project", "programme", "program", "initiative", "national", "citizen",
    "citizens", "people", "community", "development", "management",
}


def _opsi_tokens(query):
    raw = [w.lower() for w in re.findall(r"[A-Za-z][A-Za-z0-9\-]{2,}", query or "")]
    toks = [t for t in raw if t not in _OPSI_STOP]
    return toks or raw            # 전부 불용어면 원본으로 폴백


def _opsi_lookup(query):
    """영어 질의어로 OPSI 사례(제목·본문·국가)를 검색해 상위 5건. 본문(summary) 포함.
    제목 일치와 인접 구절(예: 'basic income')에 가중해 변별력을 높인다."""
    toks = _opsi_tokens(query)
    if not toks or not OPSI_DB.exists():
        return []
    try:
        conn = sqlite3.connect(OPSI_DB)
        conn.row_factory = sqlite3.Row
        where = " OR ".join(["cleaned_content LIKE ? OR title LIKE ? OR country LIKE ?"] * len(toks))
        params = []
        for t in toks:
            params += [f"%{t}%", f"%{t}%", f"%{t}%"]
        rows = [dict(r) for r in conn.execute(
            f"SELECT * FROM cases WHERE {where} LIMIT 120", params).fetchall()]
        conn.close()
    except Exception as e:
        print("opsi lookup 실패:", e, file=sys.stderr)
        return []
    scored = []
    for d in rows:
        title = (d.get("title") or "").lower()
        hay = f"{title} {(d.get('cleaned_content') or '').lower()}"
        score = 0
        for t in toks:
            if t in title:
                score += 3            # 제목 일치는 강한 신호
            elif t in hay:
                score += 1
        for a, b in zip(toks, toks[1:]):
            if f"{a} {b}" in hay:     # 인접 구절이 통째로 나오면 큰 보너스
                score += 3
        if score:
            scored.append((score, d))
    scored.sort(key=lambda x: x[0], reverse=True)
    out = []
    for _, d in scored[:5]:
        out.append({
            "title": d.get("title"), "country": d.get("country"), "year": d.get("year"),
            "sector": d.get("sector"), "level": d.get("level_of_government"),
            "summary": (d.get("cleaned_content") or "")[:500], "url": d.get("source_url"),
        })
    return out


def _originality_payload(result):
    o = result["originality"]
    spec, judge, lookup = result["spec"], result["judge"], result.get("lookup")
    return {
        "band": o["band"], "confidence": o["confidence"],
        "evidence": o["evidence"], "reasoning": o["reasoning"],
        "retraction_condition": o.get("retraction_condition", ""),
        "policy_type": spec.get("policy_type"),
        "verdict": judge.get("verdict"),
        "claimed_precedents": spec.get("claimed_precedents", []),
        "lookup": None if lookup is None else {
            "profile": lookup["profile"], "queries": lookup["queries"],
            "fiscal": lookup["fiscal"], "prism": lookup["prism"], "bill": lookup["bill"],
            "overseas": lookup.get("overseas", []),
        },
    }


def _run_originality_axis(sid):
    """백그라운드 스레드: 축 B를 돌려 결과를 DB에 저장한다.
    call_claude가 블로킹이므로 asyncio 대신 스레드로 격리한다."""
    try:
        transcript = "\n\n".join(_transcript_log(sid))
        fiscal_fn = _fiscal_local_search if _fiscal_available() else None
        # '검토(연구)' 슬롯: KDI 코퍼스가 있으면 그것, 없으면 PRISM(플래그 켜졌을 때만).
        prism_fn = (_kdi_lookup if _kdi_available()
                    else (_prism_lookup if (PRISM_ENABLED and DATA_GO_KR_KEY) else None))
        bill_fn = _bill_lookup if ASSEMBLY_KEY else None
        overseas_fn = _opsi_lookup if _opsi_available() else None
        result = engine.originality_axis(transcript, fiscal_fn, prism_fn, bill_fn, overseas_fn)
        db.save_originality(sid, result)
    except Exception as e:
        db.set_originality_status(sid, "error")
        print("originality axis 실패:", e, file=sys.stderr)


@app.post("/api/sessions/{sid}/originality")
def start_originality(sid: str):
    """축 B를 백그라운드로 시작한다(정책 프로필 전용). 즉시 pending 반환."""
    session = db.get_session(sid)
    if not session:
        raise HTTPException(404, "세션이 없습니다.")
    if session["profile"] != "정책":
        raise HTTPException(400, "정책 프로필 세션에서만 선례 조사를 실행합니다.")
    cur = db.get_originality(sid)
    if cur and cur["status"] in ("pending", "done"):
        return {"status": cur["status"]}
    db.set_originality_status(sid, "pending")
    threading.Thread(target=_run_originality_axis, args=(sid,), daemon=True).start()
    return {"status": "pending"}


@app.get("/api/sessions/{sid}/originality")
def poll_originality(sid: str):
    o = db.get_originality(sid)
    if o is None:
        raise HTTPException(404, "세션이 없습니다.")
    payload = {"status": o["status"]}
    if o["result"]:
        payload["axis_b"] = _originality_payload(o["result"])
    return payload


# ── 문서 불러오기(.docx/.pdf/.txt/.json) → 텍스트 추출 ───────────────────
def _extract_docx_bytes(data):
    from docx import Document                       # python-docx
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for tbl in doc.tables:                           # 표로 저장한 문답도 포함
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def _extract_pdf_bytes(data):
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n\n".join(pages).strip()


def _extract_document(filename, data):
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    if ext == "docx":
        return _extract_docx_bytes(data)
    if ext == "pdf":
        return _extract_pdf_bytes(data)
    if ext in ("txt", "md"):
        return data.decode("utf-8", "replace").strip()
    if ext == "json":                                # 우리 저장 형식
        obj = json.loads(data.decode("utf-8", "replace"))
        if isinstance(obj, str):
            return obj.strip()
        if isinstance(obj, dict):
            return str(obj.get("transcript_text") or obj.get("idea") or "").strip()
        return ""
    raise ValueError(f"지원하지 않는 형식입니다(.{ext}). .json/.docx/.pdf/.txt 만 됩니다.")


class ExtractRequest(BaseModel):
    filename: str
    content_b64: str


@app.post("/api/extract")
def api_extract(req: ExtractRequest):
    """업로드한 Word/PDF/텍스트/JSON에서 '문답 텍스트'를 뽑아 돌려준다."""
    if not req.filename or not req.content_b64:
        raise HTTPException(400, "파일 정보가 비어 있습니다.")
    try:
        data = base64.b64decode(req.content_b64)
    except Exception:
        raise HTTPException(400, "파일 디코딩에 실패했습니다.")
    if len(data) > 8 * 1024 * 1024:
        raise HTTPException(400, "파일이 너무 큽니다(8MB 이하만 됩니다).")
    try:
        text = _extract_document(req.filename, data)
    except ImportError:
        raise HTTPException(500, "문서 추출 라이브러리가 없습니다. 터미널에서 "
                                 "'pip install python-docx pypdf' 후 서버를 다시 시작하세요.")
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(422, f"문서에서 텍스트를 추출하지 못했습니다: {e}")
    if not text:
        raise HTTPException(422, "문서에서 텍스트를 찾지 못했습니다(스캔 이미지 PDF 등일 수 있습니다).")
    return {"text": text, "chars": len(text)}


# ── 즉석 재평가: 불러온 문답(수정본)으로 축 B만 다시 실행 ─────────────────
class AdhocRequest(BaseModel):
    transcript: str
    access_code: Optional[str] = None


@app.post("/api/originality/adhoc")
def start_originality_adhoc(req: AdhocRequest):
    """문답 전문(파일에서 불러와 수정한 것)만으로 선례 조사(축 B)를 실행한다.
    원칙(먼저 문답, 그다음 평가)에 맞게 '저장된 대화'를 대상으로만 쓴다."""
    if ACCESS_CODE and (req.access_code or "").strip() != ACCESS_CODE:
        raise HTTPException(403, "초대 코드가 올바르지 않습니다.")
    transcript = (req.transcript or "").strip()
    if not transcript:
        raise HTTPException(400, "문답 내용이 비어 있습니다.")
    if db.count_sessions_today() >= MAX_SESSIONS_PER_DAY:
        raise HTTPException(429, "오늘 사용 가능한 횟수를 모두 사용했습니다. 내일 다시 시도해 주세요.")
    sid = db.create_session(transcript[:80], dict(engine.DEFAULT_WEIGHTS), "정책")
    db.add_turn(sid, 1, "proposer", None, transcript)
    db.set_originality_status(sid, "pending")
    threading.Thread(target=_run_originality_axis, args=(sid,), daemon=True).start()
    return {"session_id": sid}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
